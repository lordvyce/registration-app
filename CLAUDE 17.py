import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkcalendar import DateEntry
import re
import json
from datetime import datetime, date, timedelta
import os
import csv
import threading
from tkinter import font
import webbrowser
import urllib.parse
import time
import subprocess
import platform

# --- Dependency Checks and Constants ---

# Try to import yagmail for email
try:
    import yagmail
    EMAIL_AVAILABLE = True
except ImportError:
    EMAIL_AVAILABLE = False

# Try to import pandas for Excel export
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# Try to import libraries for PDF and DOCX export
try:
    from fpdf import FPDF
    from docx import Document
    from docx.shared import Inches
    DOC_EXPORT_AVAILABLE = True
except ImportError:
    DOC_EXPORT_AVAILABLE = False

# --- File and Configuration Constants ---
APPOINTMENTS_FILE = 'appointments.json'
EMAIL_SETTINGS_FILE = 'email_settings.json'
REMINDER_SETTINGS_FILE = 'whatsapp_reminder_settings.json'
SENT_REMINDERS_FILE = 'sent_whatsapp_reminders.json'
REMINDER_LOG_FILE = 'whatsapp_reminder_log.txt'
PROCEDURE_TYPES = ["All", "DX", "US", "CT", "Mammo", "X-Ray", "Clinic Appointment"]
HEALTH_CARD_OPTIONS = ["Select Card", "Sagicor", "Guardian Life", "Canopy"]


class ModernCompactClinicSystem:
    def __init__(self):
        # --- CORRECTED INITIALIZATION ORDER ---
        # 1. Initialize all attributes and dictionaries first.
        self.appointments = []
        self.current_theme = "light"
        self.email_settings = {
            "enabled": True,
            "smtp_server": "smtp.gmail.com",
            "smtp_port": 587,
            "email_address": "",
            "app_password": "",
            "clinic_name": "Modern Clinic System",
            "clinic_address": "123 Medical Center Dr, Health City, HC 12345",
            "clinic_phone": "(555) 123-4567",
            "auto_send_email": True,
            "email_delay": 2
        }
        self.reminder_settings = {
            "enabled": True,
            "remind_3_days": True,
            "remind_1_day": True,
            "remind_morning": True,
            "remind_1_hour": True,
            "business_hours_start": "09:00",
            "business_hours_end": "18:00",
            "check_interval": 300,
            "auto_send_whatsapp": True,
            "whatsapp_delay": 3,
            "whatsapp_app_mode": False # New setting for app vs web
        }
        self.sent_reminders = {}
        self.reminder_thread = None
        self.reminder_running = False
        self.auto_save_active = True
        
        # 2. Set up the visual elements (themes and GUI).
        self.setup_themes()
        self.setup_gui()

        # 3. Now that the GUI exists, load data from files.
        #    This allows error toasts to be displayed correctly if loading fails.
        self.load_email_settings()
        self.load_data()
        self.load_reminder_data()
        
        # 4. Apply loaded settings to the UI
        self.update_settings_page_vars()

        # 5. Set up remaining functionalities and start background tasks.
        self.setup_keyboard_shortcuts()
        self.start_auto_save()
        self.start_reminder_system()

    def validate_appointment_data(self, apt):
        """Validate and clean appointment data, adding missing fields with defaults"""
        required_fields = {
            'id': 'Unknown_ID',
            'patient_name': 'Unknown Patient',
            'procedure': 'Unknown Procedure',
            'phone_number': 'Unknown Phone',
            'email': '',
            'health_card': '', # New field for health card
            'appointment_date': date.today().strftime('%Y-%m-%d'),
            'appointment_time': '09:00',
            'clinic_appointment_date': '', # Default for the new field
            'enable_reminders': True,
            'enable_email': True,
            'notes': '',
            'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        }
        
        # Ensure all required fields exist
        for field, default_value in required_fields.items():
            if field not in apt or apt[field] is None:
                apt[field] = default_value
        
        return apt

    def is_id_unique(self, patient_id):
        """Check if the provided patient ID is unique."""
        return not any(str(apt.get('id', '')) == str(patient_id) for apt in self.appointments)

    def save_email_settings(self):
        """Save email settings from the settings page"""
        try:
            # Update self.email_settings from UI vars before saving
            for key, var in self.email_settings_vars.items():
                self.email_settings[key] = var.get()
            
            with open(EMAIL_SETTINGS_FILE, 'w') as f:
                json.dump(self.email_settings, f, indent=4)
            self.show_toast("Email settings saved successfully!", "success")
        except Exception as e:
            self.show_toast(f"Error saving email settings: {e}", "error")

    def load_email_settings(self):
        """Load email settings from file"""
        try:
            if os.path.exists(EMAIL_SETTINGS_FILE):
                with open(EMAIL_SETTINGS_FILE, 'r') as f:
                    saved_settings = json.load(f)
                    self.email_settings.update(saved_settings)
        except Exception as e:
            self.show_toast(f"Could not load email settings: {e}", "error")

    def validate_email(self, email):
        """Validate email format"""
        if not email:
            return False
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return re.match(pattern, email) is not None

    def send_email_reminder(self, appointment, reminder_type):
        """Send email reminder to patient with enhanced private server support"""
        if not EMAIL_AVAILABLE:
            self.log_reminder_activity(
                appointment.get('patient_name', 'Unknown'),
                appointment.get('email', 'N/A'),
                f"Email library not available for {reminder_type} reminder",
                "ERROR"
            )
            return False

        email = appointment.get('email', '').strip()
        if not email or not self.validate_email(email):
            return False

        try:
            subject = self.get_email_subject(appointment, reminder_type)
            body = self.get_email_body(appointment, reminder_type)

            if self.email_settings.get("auto_send_email", True):
                # Use the same enhanced connection method for sending reminders
                smtp_server = self.email_settings["smtp_server"]
                smtp_port = int(self.email_settings["smtp_port"])
                email_addr = self.email_settings["email_address"]
                app_password = self.email_settings["app_password"]
                
                # Try the same methods as in test_email_connection
                import smtplib
                import ssl
                from email.mime.text import MIMEText
                from email.mime.multipart import MIMEMultipart
                
                # Create SSL context
                context = ssl.create_default_context()
                context.check_hostname = False
                context.verify_mode = ssl.CERT_NONE
                
                server = None
                
                try:
                    if smtp_port == 465:  # SSL
                        server = smtplib.SMTP_SSL(smtp_server, smtp_port, context=context)
                    elif smtp_port == 587:  # TLS
                        server = smtplib.SMTP(smtp_server, smtp_port)
                        server.starttls(context=context)
                    else:  # Try both
                        try:
                            server = smtplib.SMTP_SSL(smtp_server, smtp_port, context=context)
                        except:
                            server = smtplib.SMTP(smtp_server, smtp_port)
                            server.starttls(context=context)
                    
                    server.login(email_addr, app_password)
                    
                    # Compose message
                    msg = MIMEMultipart()
                    msg['From'] = email_addr
                    msg['To'] = email
                    msg['Subject'] = subject
                    msg.attach(MIMEText(body, 'plain'))
                    
                    server.send_message(msg)
                    server.quit()

                    self.log_reminder_activity(
                        appointment.get('patient_name', 'Unknown'),
                        email,
                        f"{reminder_type.replace('_', ' ').title()} email reminder sent",
                        "SENT 📧"
                    )
                    return True
                    
                except Exception as e:
                    print(f"Email reminder send failed: {e}")
                    # Try without SSL as fallback
                    try:
                        server = smtplib.SMTP(smtp_server, smtp_port)
                        server.login(email_addr, app_password)
                        
                        msg = MIMEMultipart()
                        msg['From'] = email_addr
                        msg['To'] = email
                        msg['Subject'] = subject
                        msg.attach(MIMEText(body, 'plain'))
                        
                        server.send_message(msg)
                        server.quit()
                        
                        self.log_reminder_activity(
                            appointment.get('patient_name', 'Unknown'),
                            email,
                            f"{reminder_type.replace('_', ' ').title()} email reminder sent (no SSL)",
                            "SENT 📧"
                        )
                        return True
                    except Exception as fallback_error:
                        self.log_reminder_activity(
                            appointment.get('patient_name', 'Unknown'),
                            email,
                            f"Error sending {reminder_type} email: {str(fallback_error)}",
                            "ERROR ❌"
                        )
                        return False
            else:
                self.log_reminder_activity(
                    appointment.get('patient_name', 'Unknown'),
                    email,
                    f"{reminder_type.replace('_', ' ').title()} email reminder (auto-send disabled)",
                    "LOGGED 📧"
                )
                return True

        except Exception as e:
            self.log_reminder_activity(
                appointment.get('patient_name', 'Unknown'),
                email,
                f"Error sending {reminder_type} email: {str(e)}",
                "ERROR ❌"
            )
            return False

    def get_email_subject(self, appointment, reminder_type):
        """Generate email subject based on reminder type"""
        clinic_name = self.email_settings.get("clinic_name", "Clinic")
        name = appointment.get('patient_name', 'Patient')

        subjects = {
            "3_days": f"Appointment Reminder - {name} | {clinic_name}",
            "1_day": f"Tomorrow's Appointment - {name} | {clinic_name}",
            "morning": f"Today's Appointment - {name} | {clinic_name}",
            "1_hour": f"Appointment in 1 Hour - {name} | {clinic_name}",
            "manual": f"Appointment Reminder - {name} | {clinic_name}"
        }

        return subjects.get(reminder_type, f"Appointment Reminder - {name} | {clinic_name}")

    def get_email_body(self, appointment, reminder_type):
        """Generate email body based on reminder type"""
        name = appointment.get('patient_name', 'Patient')
        procedure = appointment.get('procedure', 'N/A')
        apt_date = appointment.get('appointment_date', 'N/A')
        apt_time = appointment.get('appointment_time', '09:00')
        clinic_name = self.email_settings.get("clinic_name", "Clinic")
        clinic_address = self.email_settings.get("clinic_address", "")
        clinic_phone = self.email_settings.get("clinic_phone", "")

        bodies = {
                "3_days": f"""
    Dear {name},

    This is a friendly reminder about your upcoming appointment:

    🔬 Procedure: {procedure}
    📅 Date: {apt_date}
    ⏰ Time: {apt_time}
    📍 Location: {clinic_address}

    Your appointment is in 3 days. Please mark your calendar.

    If you need to reschedule, please contact us at {clinic_phone}.

    Best regards,
    {clinic_name} Team
            """,

            "1_day": f"""
    Dear {name},

    Your appointment is tomorrow! Here are the details:

    🔬 Procedure: {procedure}
    📅 Date: {apt_date} (TOMORROW)
    ⏰ Time: {apt_time}
    📍 Location: {clinic_address}

    Please arrive 15 minutes early for check-in.

    Contact us at {clinic_phone} if you have any questions.

    Best regards,
    {clinic_name} Team
            """,

            "morning": f"""
    Dear {name},

    Good morning! You have an appointment TODAY:

    🔬 Procedure: {procedure}
    📅 Date: TODAY ({apt_date})
    ⏰ Time: {apt_time}
    📍 Location: {clinic_address}

    Our team is ready to assist you.

    If you're running late, please call us immediately at {clinic_phone}.

    Best regards,
    {clinic_name} Team
            """,

            "1_hour": f"""
    Dear {name},

    Your appointment is in 1 HOUR:

    🔬 Procedure: {procedure}
    ⏰ Time: {apt_time} (in 1 hour)
    📍 Location: {clinic_address}

    Please make your way to our clinic.

    If you're running late, please call us at {clinic_phone}.

    Best regards,
    {clinic_name} Team
            """
        }

        return bodies.get(reminder_type, f"""
    Dear {name},

    This is a reminder about your appointment:

    🔬 Procedure: {procedure}
    📅 Date: {apt_date}
    ⏰ Time: {apt_time}
    📍 Location: {clinic_address}

    Please contact us at {clinic_phone} if you have any questions.

    Best regards,
    {clinic_name} Team
        """)

    def setup_themes(self):
        """Setup light and dark theme configurations"""
        self.themes = {
            "light": {
                "bg_primary": "#ffffff",
                "bg_secondary": "#f8fafc",
                "bg_accent": "#e2e8f0",
                "text_primary": "#1e293b",
                "text_secondary": "#64748b",
                "accent": "#3b82f6",
                "success": "#10b981",
                "warning": "#f59e0b",
                "danger": "#ef4444",
                "sidebar": "#f1f5f9",
            },
            "dark": {
                "bg_primary": "#0f172a",
                "bg_secondary": "#1e293b",
                "bg_accent": "#334155",
                "text_primary": "#f1f5f9",
                "text_secondary": "#cbd5e1",
                "accent": "#60a5fa",
                "success": "#34d399",
                "warning": "#fbbf24",
                "danger": "#f87171",
                "sidebar": "#1e293b",
            }
        }

    def get_theme(self):
        """Get current theme colors"""
        return self.themes[self.current_theme]

    def setup_gui(self):
        """Setup the main GUI"""
        self.root = tk.Tk()
        self.root.title("🏥 Modern Clinic System - Auto Reminders")
        self.root.geometry("1200x800")
        self.root.minsize(1050, 700)

        theme = self.get_theme()
        self.root.configure(bg=theme["bg_primary"])

        self.fonts = {
            "title": ("Segoe UI", 16, "bold"),
            "heading": ("Segoe UI", 12, "bold"),
            "body": ("Segoe UI", 10),
            "small": ("Segoe UI", 8)
        }

        self.setup_main_layout()
        self.create_sidebar()
        self.create_main_content()
        self.create_notification_system()

    def setup_main_layout(self):
        """Create the main layout structure"""
        self.main_container = tk.Frame(self.root)
        self.main_container.pack(fill='both', expand=True)
        self.main_container.grid_columnconfigure(1, weight=1)
        self.main_container.grid_rowconfigure(0, weight=1)

    def create_sidebar(self):
        """Create compact sidebar navigation"""
        theme = self.get_theme()
        self.sidebar = tk.Frame(
            self.main_container,
            width=200,
            bg=theme["sidebar"],
            relief='solid',
            bd=1,
            highlightthickness=1,
            highlightbackground=theme["bg_accent"]
        )
        self.sidebar.grid(row=0, column=0, sticky='ns')
        self.sidebar.grid_propagate(False)

        header_frame = tk.Frame(self.sidebar, bg=theme["sidebar"])
        header_frame.pack(fill='x', pady=10)

        title_label = tk.Label(header_frame, text="🏥 Clinic Pro", font=self.fonts["title"], bg=theme["sidebar"], fg=theme["text_primary"])
        title_label.pack()

        whatsapp_status = "📱 WhatsApp AUTO" if self.reminder_settings.get("auto_send_whatsapp", True) else "📱 Manual Only"
        self.reminder_status = tk.Label(header_frame, text=f"🔔 Reminders: {'ON' if self.reminder_settings['enabled'] else 'OFF'}", font=self.fonts["small"], bg=theme["success"] if self.reminder_settings["enabled"] else theme["danger"], fg="white", padx=5, pady=2)
        self.reminder_status.pack(pady=(5, 2))

        self.whatsapp_status = tk.Label(header_frame, text=whatsapp_status, font=self.fonts["small"], bg=theme["accent"], fg="white", padx=5, pady=2)
        self.whatsapp_status.pack(pady=(2, 0))

        self.theme_btn = tk.Button(header_frame, text="🌙" if self.current_theme == "light" else "☀️", command=self.toggle_theme, bg=theme["accent"], fg="white", font=self.fonts["body"], relief="flat", padx=10, pady=3)
        self.theme_btn.pack(pady=(10, 0))

        self.nav_buttons = []
        nav_items = [
            ("➕ Add Patient", self.show_add_page, "add"),
            ("📋 View All", self.show_view_page, "view"),
            ("📊 Dashboard & Reports", self.show_dashboard_page, "dashboard"),
            ("📱 Auto Reminders", self.show_reminders_page, "reminders"),
            ("⚙️ Settings", self.show_settings_page, "settings")
        ]

        nav_frame = tk.Frame(self.sidebar, bg=theme["sidebar"])
        nav_frame.pack(fill='x', pady=20, padx=10)
        self.current_page = "add"

        for text, command, page_id in nav_items:
            btn = tk.Button(nav_frame, text=text, command=lambda cmd=command, pid=page_id: self.navigate_to(cmd, pid), bg=theme["accent"] if page_id == self.current_page else theme["bg_accent"], fg="white" if page_id == self.current_page else theme["text_primary"], relief="flat", pady=5)
            btn.pack(fill='x', pady=2)
            self.nav_buttons.append((btn, page_id))

        self.create_quick_stats()

    def create_quick_stats(self):
        """Create quick statistics panel in sidebar"""
        theme = self.get_theme()
        stats_frame = tk.LabelFrame(self.sidebar, text="📈 Live Stats", bg=theme["sidebar"], fg=theme["text_primary"], font=self.fonts["small"], relief='flat')
        stats_frame.pack(fill='x', pady=20, padx=10)

        self.stats_labels = {}
        stats_items = [("Total Patients:", "total"), ("Today's Appts:", "today"), ("Reminders Sent:", "whatsapp_sent")]

        for text, key in stats_items:
            frame = tk.Frame(stats_frame, bg=theme["sidebar"])
            frame.pack(fill='x', pady=2)
            tk.Label(frame, text=text, bg=theme["sidebar"], fg=theme["text_secondary"], font=self.fonts["small"]).pack(side='left')
            value_label = tk.Label(frame, text="0", bg=theme["sidebar"], fg=theme["accent"], font=("Segoe UI", 8, "bold"))
            value_label.pack(side='right')
            self.stats_labels[key] = value_label

    def create_main_content(self):
        """Create main content area"""
        theme = self.get_theme()
        self.content_area = tk.Frame(self.main_container, bg=theme["bg_primary"])
        self.content_area.grid(row=0, column=1, sticky='nsew', padx=10, pady=10)

        self.pages = {}
        self.create_add_page()
        self.create_view_page()
        self.create_dashboard_page() # Combined page
        self.create_reminders_page()
        self.create_settings_page()

        self.show_add_page()

    def create_add_page(self):
        """Add patient page with a manual, required Patient ID field."""
        theme = self.get_theme()
        page = tk.Frame(self.content_area, bg=theme["bg_primary"])
        self.pages["add"] = page

        header = tk.Frame(page, bg=theme["bg_primary"])
        header.pack(fill='x', pady=(0, 20))
        tk.Label(header, text="➕ New Patient Appointment", font=self.fonts["title"], bg=theme["bg_primary"], fg=theme["text_primary"]).pack(side='left')

        form_frame = tk.Frame(page, bg=theme["bg_secondary"], relief='flat', bd=1)
        form_frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.form_vars = {}
        main_form = tk.Frame(form_frame, bg=theme["bg_secondary"])
        main_form.pack(padx=30, pady=30)
        main_form.grid_columnconfigure(1, weight=1)
        row = 0

        # --- Patient ID Field ---
        tk.Label(main_form, text="🆔 Patient ID*", font=self.fonts["heading"], bg=theme["bg_secondary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='w', pady=(5, 2))
        self.form_vars["id"] = tk.StringVar()
        tk.Entry(main_form, textvariable=self.form_vars["id"], font=self.fonts["body"], bg=theme["bg_primary"], fg=theme["text_primary"], relief='solid', bd=1, width=40).grid(row=row, column=1, sticky='ew', pady=(5, 5), padx=(10, 0))
        row += 1

        # Patient Name
        tk.Label(main_form, text="👤 Patient Name*", font=self.fonts["heading"], bg=theme["bg_secondary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='w', pady=(5, 2))
        self.form_vars["name"] = tk.StringVar()
        tk.Entry(main_form, textvariable=self.form_vars["name"], font=self.fonts["body"], bg=theme["bg_primary"], fg=theme["text_primary"], relief='solid', bd=1, width=40).grid(row=row, column=1, sticky='ew', pady=(5, 5), padx=(10, 0))
        row += 1

        # Procedure Type & Detail
        tk.Label(main_form, text="🔬 Procedure*", font=self.fonts["heading"], bg=theme["bg_secondary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='w', pady=(5, 2))
        proc_frame = tk.Frame(main_form, bg=theme["bg_secondary"])
        proc_frame.grid(row=row, column=1, sticky='ew', pady=(5, 10), padx=(10, 0))
        self.form_vars["procedure_type"] = tk.StringVar()
        proc_combo = ttk.Combobox(proc_frame, textvariable=self.form_vars["procedure_type"], values=PROCEDURE_TYPES[1:], state="readonly", width=12, font=self.fonts["body"])
        proc_combo.pack(side='left')
        proc_combo.set("Select Type")
        self.form_vars["procedure_details"] = tk.StringVar()
        details_entry = tk.Entry(proc_frame, textvariable=self.form_vars["procedure_details"], font=self.fonts["body"], bg=theme["bg_primary"], fg=theme["text_primary"], relief='solid', bd=1)
        details_entry.pack(side='left', fill='x', expand=True, padx=(8, 0))
        row += 1

        # WhatsApp Phone Number
        tk.Label(main_form, text="📱 WhatsApp Phone*", font=self.fonts["heading"], bg=theme["bg_secondary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='w', pady=(5, 2))
        self.form_vars["phone1"] = tk.StringVar()
        phone_entry = tk.Entry(main_form, textvariable=self.form_vars["phone1"], font=self.fonts["body"], bg=theme["bg_primary"], fg=theme["text_primary"], relief='solid', bd=1)
        phone_entry.grid(row=row, column=1, sticky='ew', pady=(5, 5), padx=(10, 0))
        row += 1
        tk.Label(main_form, text="💡 Include country code (e.g., +1234567890)", font=self.fonts["small"], bg=theme["bg_secondary"], fg=theme["text_secondary"], anchor='w').grid(row=row, column=1, sticky='w', pady=(0, 5), padx=(10, 0))
        row += 1

        # Email Address
        tk.Label(main_form, text="📧 Email (Optional)", font=self.fonts["body"], bg=theme["bg_secondary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='w', pady=(5, 2))
        self.form_vars["email"] = tk.StringVar()
        tk.Entry(main_form, textvariable=self.form_vars["email"], font=self.fonts["body"], bg=theme["bg_primary"], fg=theme["text_primary"], relief='solid', bd=1).grid(row=row, column=1, sticky='ew', pady=(5, 5), padx=(10, 0))
        row += 1

        # --- NEW Health Card Dropdown ---
        tk.Label(main_form, text="💳 Health Card (Optional)", font=self.fonts["body"], bg=theme["bg_secondary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='w', pady=(5, 2))
        self.form_vars["health_card"] = tk.StringVar()
        health_card_combo = ttk.Combobox(main_form, textvariable=self.form_vars["health_card"], values=HEALTH_CARD_OPTIONS, state="readonly", font=self.fonts["body"])
        health_card_combo.grid(row=row, column=1, sticky='ew', pady=(5, 10), padx=(10, 0))
        health_card_combo.set(HEALTH_CARD_OPTIONS[0]) # Set default value
        row += 1
        # --- END NEW SECTION ---

        # Appointment Date & Time
        tk.Label(main_form, text="📅 Appt. Date & Time*", font=self.fonts["heading"], bg=theme["bg_secondary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='w', pady=(5, 2))
        datetime_frame = tk.Frame(main_form, bg=theme["bg_secondary"])
        datetime_frame.grid(row=row, column=1, sticky='ew', pady=(5, 10), padx=(10, 0))
        self.appointment_date_entry = DateEntry(datetime_frame, width=15, background=theme["accent"], foreground='white', borderwidth=1, font=self.fonts["body"], mindate=date.today())
        self.appointment_date_entry.pack(side='left')
        self.form_vars["appointment_time"] = tk.StringVar(value="09:00")
        time_entry = tk.Entry(datetime_frame, textvariable=self.form_vars["appointment_time"], font=self.fonts["body"], bg=theme["bg_primary"], fg=theme["text_primary"], relief='solid', bd=1, width=8)
        time_entry.pack(side='left', padx=(8, 0))
        tk.Label(datetime_frame, text="(HH:MM 24hr)", font=self.fonts["small"], bg=theme["bg_secondary"], fg=theme["text_secondary"]).pack(side='left', padx=(5, 0))
        row += 1

        # --- NEW SECTION: Clinic Appointment Date ---
        tk.Label(main_form, text="🏥 Clinic Appt. Date (Optional)", font=self.fonts["body"], bg=theme["bg_secondary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='w', pady=(5, 2))
        self.clinic_appointment_date_entry = DateEntry(main_form, width=15, background=theme["accent"], foreground='white', borderwidth=1, font=self.fonts["body"], mindate=date.today())
        self.clinic_appointment_date_entry.grid(row=row, column=1, sticky='ew', pady=(5, 10), padx=(10, 0))
        self.clinic_appointment_date_entry.set_date(None) # Optional: Set to None initially
        row += 1
        # --- END NEW SECTION ---

        # Communication Preferences
        communication_frame = tk.LabelFrame(main_form, text="📱📧 Auto-Reminders", bg=theme["bg_secondary"], fg=theme["accent"], font=self.fonts["heading"])
        communication_frame.grid(row=row, column=0, columnspan=2, sticky='ew', pady=15)
        self.form_vars["enable_reminders"] = tk.BooleanVar(value=True)
        tk.Checkbutton(communication_frame, text="Send automatic WhatsApp reminders", variable=self.form_vars["enable_reminders"], bg=theme["bg_secondary"], fg=theme["text_primary"], font=self.fonts["body"]).pack(anchor='w', padx=15, pady=5)
        self.form_vars["enable_email"] = tk.BooleanVar(value=True)
        tk.Checkbutton(communication_frame, text="Send automatic email reminders", variable=self.form_vars["enable_email"], bg=theme["bg_secondary"], fg=theme["text_primary"], font=self.fonts["body"]).pack(anchor='w', padx=15, pady=5)
        row += 1

        # Notes
        tk.Label(main_form, text="📝 Notes (Optional)", font=self.fonts["heading"], bg=theme["bg_secondary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='nw', pady=(5, 2))
        self.notes_text = tk.Text(main_form, font=self.fonts["body"], bg=theme["bg_primary"], fg=theme["text_primary"], relief='solid', bd=1, height=2, width=40, wrap='word')
        self.notes_text.grid(row=row, column=1, sticky='ew', pady=(5, 15), padx=(10, 0))
        row += 1

        # Action buttons
        actions_frame = tk.Frame(main_form, bg=theme["bg_secondary"])
        actions_frame.grid(row=row, column=0, columnspan=2, pady=20)
        tk.Button(actions_frame, text="💾 Save Patient", command=self.add_appointment, bg=theme["success"], fg="white", font=self.fonts["heading"], relief='flat', padx=20, pady=8, cursor='hand2').pack(side='left', padx=10)
        tk.Button(actions_frame, text="🗑️ Clear Form", command=self.clear_form, bg=theme["warning"], fg="white", font=self.fonts["body"], relief='flat', padx=15, pady=6, cursor='hand2').pack(side='left')

    def create_filter_controls(self, parent_frame):
        """Creates a reusable frame with filter widgets."""
        theme = self.get_theme()
        filter_frame = tk.LabelFrame(parent_frame, text="🔍 Filter Appointments", bg=theme["bg_secondary"], fg=theme["text_primary"], font=self.fonts["body"])
        filter_frame.pack(fill='x', padx=5, pady=(0, 10))

        controls_frame = tk.Frame(filter_frame, bg=theme["bg_secondary"])
        controls_frame.pack(fill='x', padx=10, pady=10)

        # --- New: Patient ID Filter ---
        tk.Label(controls_frame, text="ID:", bg=theme["bg_secondary"], fg=theme["text_primary"]).grid(row=0, column=0, padx=(0, 5))
        id_var = tk.StringVar()
        tk.Entry(controls_frame, textvariable=id_var, width=10).grid(row=0, column=1, padx=(0, 15))

        tk.Label(controls_frame, text="Name:", bg=theme["bg_secondary"], fg=theme["text_primary"]).grid(row=0, column=2, padx=(0, 5)) # Adjusted column
        name_var = tk.StringVar()
        tk.Entry(controls_frame, textvariable=name_var, width=15).grid(row=0, column=3, padx=(0, 15)) # Adjusted column

        tk.Label(controls_frame, text="From:", bg=theme["bg_secondary"], fg=theme["text_primary"]).grid(row=0, column=4, padx=(0, 5)) # Adjusted column
        start_date_entry = DateEntry(controls_frame, width=12, background=theme["accent"], foreground='white', borderwidth=1)
        start_date_entry.set_date(None)
        start_date_entry.grid(row=0, column=5, padx=(0, 10)) # Adjusted column

        tk.Label(controls_frame, text="To:", bg=theme["bg_secondary"], fg=theme["text_primary"]).grid(row=0, column=6, padx=(0, 5)) # Adjusted column
        end_date_entry = DateEntry(controls_frame, width=12, background=theme["accent"], foreground='white', borderwidth=1)
        end_date_entry.set_date(None)
        end_date_entry.grid(row=0, column=7, padx=(0, 15)) # Adjusted column

        tk.Label(controls_frame, text="Modality:", bg=theme["bg_secondary"], fg=theme["text_primary"]).grid(row=0, column=8, padx=(0, 5)) # Adjusted column
        modality_var = tk.StringVar(value="All")
        ttk.Combobox(controls_frame, textvariable=modality_var, values=PROCEDURE_TYPES, state="readonly", width=10).grid(row=0, column=9, padx=(0, 15)) # Adjusted column

        filter_widgets = { "id": id_var, "name": name_var, "start_date": start_date_entry, "end_date": end_date_entry, "modality": modality_var, "parent_frame": parent_frame } # Added "id"

        tk.Button(controls_frame, text="Apply", command=lambda: self.apply_filters(filter_widgets), bg=theme["success"], fg="white", relief='flat').grid(row=0, column=10, padx=5) # Adjusted column
        tk.Button(controls_frame, text="Clear", command=lambda: self.clear_filters(filter_widgets), bg=theme["warning"], fg="white", relief='flat').grid(row=0, column=11, padx=5) # Adjusted column

        return filter_widgets

    def apply_filters(self, widgets):
        """Applies filters to the appointments list and updates the appropriate tree."""
        id_query = widgets["id"].get().lower().strip() # New: Get ID query
        name_query = widgets["name"].get().lower().strip()
        modality_query = widgets["modality"].get()
        try:
            start_date = widgets["start_date"].get_date()
        except (tk.TclError, TypeError):
            start_date = None
        try:
            end_date = widgets["end_date"].get_date()
        except (tk.TclError, TypeError):
            end_date = None

        filtered_appointments = []
        for apt in self.appointments:
            match = True
            # New: Filter by ID
            if id_query and id_query not in str(apt.get('id', '')).lower():
                match = False
            if name_query and name_query not in apt.get('patient_name', '').lower():
                match = False
            if modality_query != "All" and not apt.get('procedure', '').upper().startswith(modality_query):
                match = False
            if start_date and end_date:
                try:
                    apt_date = datetime.strptime(apt.get('appointment_date', ''), '%Y-%m-%d').date()
                    if not (start_date <= apt_date <= end_date):
                        match = False
                except (ValueError, KeyError):
                    match = False
            if match:
                filtered_appointments.append(apt)

        if widgets["parent_frame"] == self.pages["view"]:
            self.refresh_appointments(filtered_appointments)
        elif widgets["parent_frame"] == self.pages["dashboard"]:
            self.update_dashboard_page(filtered_appointments)

    def clear_filters(self, widgets):
        """Clears filter widgets and refreshes the view."""
        widgets["id"].set("") # New: Clear ID filter
        widgets["name"].set("")
        widgets["start_date"].set_date(None)
        widgets["end_date"].set_date(None)
        widgets["modality"].set("All")
        if widgets["parent_frame"] == self.pages["view"]:
            self.refresh_appointments(self.appointments)
        elif widgets["parent_frame"] == self.pages["dashboard"]:
             self.update_dashboard_page(self.appointments)

    def create_view_page(self):
        """Create compact appointments view page with filtering and export."""
        theme = self.get_theme()
        page = tk.Frame(self.content_area, bg=theme["bg_primary"])
        self.pages["view"] = page
        header = tk.Frame(page, bg=theme["bg_primary"])
        header.pack(fill='x', pady=(0, 10))
        tk.Label(header, text="📋 All Appointments", font=self.fonts["title"], bg=theme["bg_primary"], fg=theme["text_primary"]).pack(side='left')

        self.view_filters = self.create_filter_controls(page)
        tree_frame = tk.Frame(page, bg=theme["bg_secondary"], relief='flat', bd=1)
        tree_frame.pack(fill='both', expand=True)

        controls = tk.Frame(tree_frame, bg=theme["bg_secondary"])
        controls.pack(fill='x', pady=5, padx=5)
        tk.Button(controls, text="🗑️ Delete", command=self.delete_appointment, bg=theme["danger"], fg="white", relief='flat', padx=10).pack(side='left', padx=2)
        
        # New button for manual email
        tk.Button(controls, text="📧 Manual Email", command=self.send_manual_email, bg=theme["accent"], fg="white", relief='flat', padx=10).pack(side='left', padx=2)

        tk.Button(controls, text="📱 Manual WhatsApp", command=self.send_manual_whatsapp, bg=theme["success"], fg="white", relief='flat', padx=10).pack(side='left', padx=2)
        tk.Button(controls, text="📄 Export Schedule", command=self.create_export_dialog, bg=theme["warning"], fg="black", relief='flat', padx=10).pack(side='right', padx=2)

        v_scrollbar = ttk.Scrollbar(tree_frame)
        v_scrollbar.pack(side='right', fill='y')
        h_scrollbar = ttk.Scrollbar(tree_frame, orient='horizontal')
        h_scrollbar.pack(side='bottom', fill='x')
        
        # Configure style for the treeview
        style = ttk.Style()
        style.configure("Treeview")
        style.map("Treeview",
              background=[('selected', theme["accent"])],
              foreground=[('selected', "white")])

        # Modified: Updated columns to separate Modality and Procedure
        self.appointments_tree = ttk.Treeview(tree_frame, columns=('ID', 'Name', 'Modality', 'Procedure', 'Phone', 'Email', 'Health Card', 'DateTime', 'ClinicApptDate', 'Notes'), show='headings', yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        v_scrollbar.config(command=self.appointments_tree.yview)
        h_scrollbar.config(command=self.appointments_tree.xview)
        
        # Configure tags directly on the treeview widget
        self.appointments_tree.tag_configure('GuardianLife', foreground='purple')
        self.appointments_tree.tag_configure('Sagicor', foreground='sky blue') 
        self.appointments_tree.tag_configure('Canopy', foreground='lime green')

        # Modified: Added Modality column, adjusted widths
        columns = [('ID', 70), ('Name', 120), ('Modality', 80), ('Procedure', 150), ('Phone', 100), ('Email', 120), ('Health Card', 100), ('DateTime', 100), ('ClinicApptDate', 100), ('Notes', 150)]
        for col, width in columns:
            self.appointments_tree.heading(col, text=col, command=lambda c=col: self.sort_treeview(self.appointments_tree, c, False), anchor='center')
            self.appointments_tree.column(col, width=width, anchor='center')
        self.appointments_tree.pack(fill='both', expand=True, side='bottom')

        # Bind double-click event to edit_appointment
        self.appointments_tree.bind("<Double-1>", self.on_tree_double_click)

    def on_tree_double_click(self, event):
        """Handle double-click on treeview item to edit."""
        item_id = self.appointments_tree.identify_row(event.y)
        if item_id:
            self.appointments_tree.selection_set(item_id) # Select the row that was double-clicked
            self.edit_appointment()

    def create_dashboard_page(self):
        """Creates the combined Dashboard & Reports page."""
        theme = self.get_theme()
        page = tk.Frame(self.content_area, bg=theme["bg_primary"])
        self.pages["dashboard"] = page

        # Page Header
        tk.Label(page, text="📊 Dashboard & Reports", font=self.fonts["title"], bg=theme["bg_primary"], fg=theme["text_primary"]).pack(pady=(0, 10), anchor='w')

        # Top Row: Stat Cards (reflecting the entire dataset)
        stats_container = tk.Frame(page, bg=theme["bg_primary"])
        stats_container.pack(fill='x', pady=(0, 10))
        self.create_stat_card(stats_container, "Total Appointments", "📋", "total_count")
        self.create_stat_card(stats_container, "Reminders Sent", "📱", "whatsapp_count")
        self.create_stat_card(stats_container, "Appts This Week", "📅", "week_count")

        # Middle Row: Filter Controls
        self.dashboard_filters = self.create_filter_controls(page)

        # Bottom Area: Split View (Reports Tree on Left, Procedure Chart on Right)
        bottom_container = tk.Frame(page, bg=theme["bg_primary"])
        bottom_container.pack(fill='both', expand=True, pady=(10, 0))

        # Left side: Filtered Results Treeview
        results_frame = tk.LabelFrame(bottom_container, text="Filtered Appointment List", bg=theme["bg_secondary"], fg=theme["text_primary"], font=self.fonts["heading"], relief='flat', bd=1)
        results_frame.pack(side='left', fill='both', expand=True, padx=(0, 10))
        
        results_scrollbar = ttk.Scrollbar(results_frame)
        results_scrollbar.pack(side='right', fill='y')
        
        # Modified: Updated dashboard tree to match main view with separated Modality and Procedure
        self.dashboard_tree = ttk.Treeview(results_frame, columns=('ID', 'Name', 'Modality', 'Procedure', 'Phone', 'Health Card', 'DateTime', 'ClinicApptDate'), show='headings', yscrollcommand=results_scrollbar.set)
        results_scrollbar.config(command=self.dashboard_tree.yview)

        # Modified: Added Modality column
        for col, width in [('ID', 70), ('Name', 120), ('Modality', 80), ('Procedure', 150), ('Phone', 100), ('Health Card', 100), ('DateTime', 100), ('ClinicApptDate', 100)]:
            self.dashboard_tree.heading(col, text=col, command=lambda c=col: self.sort_treeview(self.dashboard_tree, c, False))
            self.dashboard_tree.column(col, width=width, anchor='w')
        self.dashboard_tree.pack(fill='both', expand=True)

        # Right side: Procedure Distribution (dynamic based on filter)
        procedures_frame = tk.LabelFrame(bottom_container, text="🔬 Procedure Distribution", bg=theme["bg_secondary"], fg=theme["text_primary"], font=self.fonts["heading"], relief='flat', bd=1)
        procedures_frame.pack(side='right', fill='both', expand=True, ipadx=10, ipady=5)
        
        self.procedure_listbox = tk.Listbox(procedures_frame, bg=theme["bg_primary"], fg=theme["text_primary"], font=self.fonts["body"], relief='flat')
        self.procedure_listbox.pack(fill='both', expand=True, padx=10, pady=10)

    def update_dashboard_page(self, filtered_data=None):
        """Updates the entire Dashboard & Reports page based on filtered or all data."""
        # If no filtered data is provided (e.g., on initial load), use all appointments
        if filtered_data is None:
            filtered_data = self.appointments
        
        # --- PART 1: Update top STAT CARDS (always based on FULL dataset) ---
        if hasattr(self, 'total_count_label'):
            self.total_count_label.config(text=str(len(self.appointments)))
        if hasattr(self, 'whatsapp_count_label'):
            self.whatsapp_count_label.config(text=str(len(self.sent_reminders)))
        
        if hasattr(self, 'week_count_label'):
            today = date.today()
            start_of_week = today - timedelta(days=today.weekday())
            end_of_week = start_of_week + timedelta(days=6)
            week_count = 0
            for apt in self.appointments:
                try:
                    apt_date = datetime.strptime(apt.get('appointment_date', ''), '%Y-%m-%d').date()
                    if start_of_week <= apt_date <= end_of_week:
                        week_count += 1
                except (ValueError, KeyError):
                    continue
            self.week_count_label.config(text=str(week_count))

        # --- PART 2: Update the appointments TREEVIEW (based on FILTERED data) ---
        if hasattr(self, 'dashboard_tree'):
            for item in self.dashboard_tree.get_children():
                self.dashboard_tree.delete(item)
            for apt in filtered_data:
                # Parse procedure to separate modality and details
                modality, procedure_details = self.parse_procedure(apt.get('procedure', 'N/A'))
                
                datetime_str = f"{apt.get('appointment_date', 'N/A')} {apt.get('appointment_time', '09:00')}"
                # Modified: Added separated modality in dashboard tree
                self.dashboard_tree.insert('', 'end', values=(
                    apt.get('id', 'N/A'), 
                    apt.get('patient_name', 'Unknown'), 
                    modality,
                    procedure_details, 
                    apt.get('phone_number', 'N/A'), 
                    apt.get('health_card', 'N/A'),
                    datetime_str,
                    apt.get('clinic_appointment_date', 'N/A')
                ))

        # --- PART 3: Update the procedure LISTBOX (based on FILTERED data) ---
        if hasattr(self, 'procedure_listbox'):
            procedure_counts = {}
            for apt in filtered_data: # Use the filtered data for dynamic analysis
                modality, _ = self.parse_procedure(apt.get('procedure', 'Unknown'))
                procedure_counts[modality] = procedure_counts.get(modality, 0) + 1
            
            self.procedure_listbox.delete(0, tk.END)
            # Sort by count descending for a clear view of top procedures
            for proc, count in sorted(procedure_counts.items(), key=lambda item: item[1], reverse=True):
                self.procedure_listbox.insert(tk.END, f"{proc}: {count} appointments")

    def parse_procedure(self, procedure_text):
        """Parse the procedure string to separate modality and details."""
        if not procedure_text or procedure_text == 'N/A':
            return 'N/A', 'N/A'
            
        # Check if procedure follows the "MODALITY: details" format
        if ':' in procedure_text:
            parts = procedure_text.split(':', 1)
            modality = parts[0].strip().upper()
            details = parts[1].strip()
            return modality, details
        else:
            # If no colon, assume the whole text is the modality
            return procedure_text.strip().upper(), ''

    def create_export_dialog(self):
        """Create a dialog to select date and format for schedule export."""
        if not DOC_EXPORT_AVAILABLE:
            self.show_toast("Export requires fpdf2 and python-docx. Please install them.", "error", 6000)
            return

        theme = self.get_theme()
        dialog = tk.Toplevel(self.root)
        dialog.title("📄 Export Daily Schedule")
        dialog.geometry("350x250")
        dialog.configure(bg=theme["bg_primary"])
        dialog.grab_set()
        dialog.resizable(False, False)

        tk.Label(dialog, text="Select Date:", bg=theme["bg_primary"], fg=theme["text_primary"], font=self.fonts["heading"]).pack(pady=(10,5))
        export_date_entry = DateEntry(dialog, width=15, background=theme["accent"], foreground='white', borderwidth=1, font=self.fonts["body"])
        export_date_entry.pack(pady=5)

        tk.Label(dialog, text="Select Format:", bg=theme["bg_primary"], fg=theme["text_primary"], font=self.fonts["heading"]).pack(pady=(10,5))
        format_var = tk.StringVar(value="PDF")
        tk.Radiobutton(dialog, text="PDF", variable=format_var, value="PDF", bg=theme["bg_primary"], fg=theme["text_primary"]).pack()
        tk.Radiobutton(dialog, text="DOCX (Word)", variable=format_var, value="DOCX", bg=theme["bg_primary"], fg=theme["text_primary"]).pack()

        def start_export():
            try:
                schedule_date = export_date_entry.get_date()
                export_format = format_var.get()
                dialog.destroy()
                self.generate_schedule_document(schedule_date, export_format)
            except (tk.TclError, TypeError):
                self.show_toast("Please select a valid date.", "error")

        tk.Button(dialog, text="Export", command=start_export, bg=theme["success"], fg="white", font=self.fonts["body"]).pack(pady=20)

    def generate_schedule_document(self, schedule_date, export_format):
        """Filters appointments for a date and generates a DOCX or PDF file."""
        date_str = schedule_date.strftime('%Y-%m-%d')
        daily_appts = [apt for apt in self.appointments if apt.get('appointment_date') == date_str]
        if not daily_appts:
            self.show_toast(f"No appointments found for {date_str}.", "warning")
            return
        daily_appts.sort(key=lambda x: datetime.strptime(x.get('appointment_time', '00:00'), '%H:%M'))

        file_extension = ".pdf" if export_format == "PDF" else ".docx"
        filename = filedialog.asksaveasfilename(defaultextension=file_extension, filetypes=[(f"{export_format} files", f"*{file_extension}"), ("All files", "*.*")], initialfile=f"Daily_Schedule_{date_str}")
        if not filename:
            return
        try:
            if export_format == "PDF":
                self.generate_pdf_schedule(filename, date_str, daily_appts)
            else: # DOCX
                self.generate_docx_schedule(filename, date_str, daily_appts)
            self.show_toast(f"Schedule exported to {export_format} successfully!", "success")
        except Exception as e:
            self.show_toast(f"Export failed: {e}", "error", 6000)

    def generate_pdf_schedule(self, filename, date_str, appointments):
        """Generates the schedule as a PDF file."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        clinic_name = self.email_settings.get("clinic_name", "Clinic")
        pdf.cell(0, 10, f"{clinic_name} - Daily Schedule", 0, 1, 'C')
        pdf.set_font("Arial", '', 12)
        pdf.cell(0, 10, f"Date: {date_str}", 0, 1, 'C')
        pdf.ln(10)

        pdf.set_font("Arial", 'B', 10)
        # Modified: Split Modality and Procedure in PDF export
        pdf.cell(25, 10, 'Time', 1)
        pdf.cell(50, 10, 'Patient Name', 1)
        pdf.cell(30, 10, 'Health Card', 1)
        pdf.cell(30, 10, 'Modality', 1) 
        pdf.cell(35, 10, 'Procedure', 1)
        pdf.cell(30, 10, 'Phone Number', 1)
        pdf.ln()

        pdf.set_font("Arial", '', 10)
        for apt in appointments:
            # Parse procedure to get modality and details
            modality, procedure_details = self.parse_procedure(apt.get('procedure', 'N/A'))
            
            pdf.cell(25, 10, apt.get('appointment_time', ''), 1)
            # Use multi_cell for potentially long names/procedures
            x_before, y_before = pdf.get_x(), pdf.get_y()
            pdf.multi_cell(50, 10, apt.get('patient_name', 'Unknown'), 1, 'L')
            pdf.set_xy(x_before + 50, y_before)
            pdf.multi_cell(30, 10, apt.get('health_card', 'N/A'), 1, 'L')
            pdf.set_xy(x_before + 80, y_before)
            pdf.multi_cell(30, 10, modality, 1, 'L') # Modality column
            pdf.set_xy(x_before + 110, y_before)
            pdf.multi_cell(35, 10, procedure_details, 1, 'L') # Procedure details column
            pdf.set_xy(x_before + 145, y_before)
            pdf.cell(30, 10, apt.get('phone_number', 'N/A'), 1)
            pdf.ln()
        pdf.output(filename)

    def generate_docx_schedule(self, filename, date_str, appointments):
        """Generates the schedule as a DOCX file."""
        doc = Document()
        clinic_name = self.email_settings.get("clinic_name", "Clinic")
        doc.add_heading(f"{clinic_name} - Daily Schedule", 0)
        doc.add_paragraph(f"Date: {date_str}", style='Intense Quote')
        # Modified: Added separated modality in DOCX export
        table = doc.add_table(rows=1, cols=6) 
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text, hdr_cells[4].text, hdr_cells[5].text = 'Time', 'Patient Name', 'Health Card', 'Modality', 'Procedure', 'Phone Number'
        for apt in appointments:
            # Parse procedure to get modality and details
            modality, procedure_details = self.parse_procedure(apt.get('procedure', 'N/A'))
            
            row_cells = table.add_row().cells
            row_cells[0].text = apt.get('appointment_time', '')
            row_cells[1].text = apt.get('patient_name', 'Unknown')
            row_cells[2].text = apt.get('health_card', 'N/A')
            row_cells[3].text = modality
            row_cells[4].text = procedure_details
            row_cells[5].text = apt.get('phone_number', 'N/A')
        doc.save(filename)


    def add_appointment(self):
            """Add new appointment using a manual, unique ID."""
            patient_id = self.form_vars.get('id', tk.StringVar()).get().strip()
            name = self.form_vars.get('name', tk.StringVar()).get().strip()
            procedure_type = self.form_vars.get('procedure_type', tk.StringVar()).get().strip()
            procedure_details = self.form_vars.get('procedure_details', tk.StringVar()).get().strip()
            phone1 = self.form_vars.get('phone1', tk.StringVar()).get().strip()
            appointment_time = self.form_vars.get('appointment_time', tk.StringVar()).get().strip()
            enable_reminders = self.form_vars.get('enable_reminders', tk.BooleanVar()).get()
            enable_email = self.form_vars.get('enable_email', tk.BooleanVar()).get() # Get email preference from form
            email = self.form_vars.get('email', tk.StringVar()).get().strip()
            notes = self.notes_text.get('1.0', tk.END).strip()
            
            # Get health card value, handle default "Select Card"
            health_card = self.form_vars.get('health_card', tk.StringVar()).get()
            if health_card == HEALTH_CARD_OPTIONS[0]:
                health_card = ""

            # Get the new clinic appointment date
            clinic_appointment_date_obj = self.clinic_appointment_date_entry.get_date()
            # Check if the date entry widget has a valid date selected, if not, set to empty string
            clinic_appointment_date_str = clinic_appointment_date_obj.strftime('%Y-%m-%d') if clinic_appointment_date_obj else ''

            if not all([patient_id, name, procedure_type, phone1]) or procedure_type == "Select Type":
                self.show_toast("Please fill all required fields (*)", "error")
                return False
            if not self.is_id_unique(patient_id):
                self.show_toast(f"Patient ID '{patient_id}' already exists. Please use a unique ID.", "error", duration=6000)
                return False
            if email and not self.validate_email(email):
                self.show_toast("Please enter a valid email address!", "error")
                return False
            if not self.clean_phone_number(phone1):
                self.show_toast("Please enter a valid WhatsApp phone number with country code!", "error")
                return False
            if not self.validate_time(appointment_time):
                self.show_toast("Please enter a valid time (HH:MM format)!", "error")
                return False

            full_procedure = f"{procedure_type.upper()}: {procedure_details}" if procedure_details else procedure_type.upper()
            appointment = {
                'id': patient_id,
                'patient_name': name,
                'procedure': full_procedure,
                'phone_number': phone1,
                'email': email,
                'health_card': health_card, # Add this new field
                'appointment_date': self.appointment_date_entry.get_date().strftime('%Y-%m-%d'),
                'appointment_time': appointment_time,
                'clinic_appointment_date': clinic_appointment_date_str, # Add this new field
                'enable_reminders': enable_reminders,
                'enable_email': enable_email,
                'notes': notes,
                'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            }
            try:
                self.appointments.append(appointment)
                self.save_data()
                self.update_stats()
                self.refresh_all()
                if enable_reminders:
                    self.log_reminder_activity(name, phone1, "New appointment created with reminders", "CREATED ✅")
                self.show_toast(f"✅ Patient '{name}' saved! Form cleared.", "success")
                self.clear_form()
                return True
            except Exception as e:
                self.show_toast(f"❌ Failed to save patient: {str(e)}", "error")
                return False

    def validate_time(self, time_str):
        """Validate time format HH:MM"""
        try:
            datetime.strptime(time_str, '%H:%M')
            return True
        except ValueError:
            return False

    def clear_form(self):
        """Clear all form fields"""
        for var in self.form_vars.values():
            if isinstance(var, tk.StringVar):
                var.set("")
            elif isinstance(var, tk.BooleanVar):
                var.set(True)
        if hasattr(self, 'form_vars') and 'procedure_type' in self.form_vars:
            self.form_vars['procedure_type'].set("Select Type")
        if hasattr(self, 'form_vars') and 'health_card' in self.form_vars:
            self.form_vars['health_card'].set(HEALTH_CARD_OPTIONS[0])
        if hasattr(self, 'appointment_date_entry'):
            self.appointment_date_entry.set_date(date.today())
        if hasattr(self, 'form_vars') and 'appointment_time' in self.form_vars:
            self.form_vars['appointment_time'].set("09:00")
        if hasattr(self, 'clinic_appointment_date_entry'): # Clear the new date entry
            self.clinic_appointment_date_entry.set_date(None)
        if hasattr(self, 'notes_text'):
            self.notes_text.delete('1.0', tk.END)

    def navigate_to(self, command, page_id):
        """Navigate to a specific page"""
        theme = self.get_theme()
        for btn, pid in self.nav_buttons:
            btn.config(bg=theme["accent"] if pid == page_id else theme["bg_accent"], fg="white" if pid == page_id else theme["text_primary"])
        self.current_page = page_id
        command()

    def show_add_page(self):
        """Show add appointment page"""
        self.hide_all_pages()
        self.pages["add"].pack(fill='both', expand=True)

    def show_reminders_page(self):
        """Show reminders page"""
        self.hide_all_pages()
        self.pages["reminders"].pack(fill='both', expand=True)
        self.refresh_reminder_log()

    def hide_all_pages(self):
        """Hide all pages"""
        for page in self.pages.values():
            page.pack_forget()

    def show_toast(self, message, mtype="info", duration=4000):
        """Show toast notification"""
        theme = self.get_theme()
        colors = { "info": theme["accent"], "success": theme["success"], "warning": theme["warning"], "error": theme["danger"] }
        toast = tk.Toplevel(self.root)
        toast.wm_overrideredirect(True)
        toast.wm_attributes("-topmost", True)
        toast.geometry(f"350x70+{self.root.winfo_x() + self.root.winfo_width() - 370}+{self.root.winfo_y() + 40}")
        frame = tk.Frame(toast, bg=colors[mtype], relief='solid', bd=1)
        frame.pack(fill='both', expand=True)
        tk.Label(frame, text=message, bg=colors[mtype], fg="white", font=self.fonts["body"], wraplength=320, justify='center').pack(expand=True, padx=15, pady=15)
        self.root.after(duration, toast.destroy)

    def refresh_appointments(self, data_source=None):
        """Refresh appointments display"""
        if data_source is None:
            data_source = self.appointments

        for item in self.appointments_tree.get_children():
            self.appointments_tree.delete(item)

        for apt in data_source:
            datetime_str = f"{apt.get('appointment_date', 'N/A')} {apt.get('appointment_time', '09:00')}"
            
            # Get health card value for tag-based coloring
            health_card = apt.get('health_card', '')
            tag = ''
            if health_card == 'Guardian Life':
                tag = 'GuardianLife'
            elif health_card == 'Sagicor':
                tag = 'Sagicor'
            elif health_card == 'Canopy':
                tag = 'Canopy'
                
            # Parse procedure to separate modality and details
            modality, procedure_details = self.parse_procedure(apt.get('procedure', 'N/A'))
            
            # Insert data into the treeview with modality and procedure separated
            self.appointments_tree.insert('', 'end', values=(
                apt.get('id', 'N/A'), 
                apt.get('patient_name', 'Unknown'), 
                modality,  # Now displays just the modality (e.g., CT, US, etc.)
                procedure_details,  # Now displays just the procedure details
                apt.get('phone_number', 'N/A'), 
                apt.get('email', ''), 
                health_card,
                datetime_str, 
                apt.get('clinic_appointment_date', 'N/A'),
                apt.get('notes', ''),
            ), tags=(tag,))
            
        self.update_stats()
        def sort_treeview(self, tree, col, reverse):
            """Sort treeview contents when a column header is clicked."""
            data = [(tree.set(child, col), child) for child in tree.get_children('')]
        try:
            # Try to sort as numbers if possible
            data.sort(key=lambda t: float(t[0]), reverse=reverse)
        except ValueError:
            # Fallback to string sorting
            data.sort(reverse=reverse)
        for index, (val, child) in enumerate(data):
            tree.move(child, '', index)
        tree.heading(col, command=lambda: self.sort_treeview(tree, col, not reverse))

    def quick_save_and_new(self):
        """Quick save and prepare for new entry"""
        if self.add_appointment():
            self.show_toast("Ready for next entry.", "info")

    def show_view_page(self):
        """Show view appointments page and refresh its content."""
        self.hide_all_pages()
        self.pages["view"].pack(fill='both', expand=True)
        self.refresh_appointments(self.appointments)
        self.clear_filters(self.view_filters)

    def show_dashboard_page(self):
        """Show combined Dashboard & Reports page."""
        self.hide_all_pages()
        self.pages["dashboard"].pack(fill='both', expand=True)
        # When showing the page, update with full data and clear filters
        self.update_dashboard_page(self.appointments) 
        if hasattr(self, 'dashboard_filters'):
            self.clear_filters(self.dashboard_filters)

    def refresh_all(self):
        """Refresh all data and views."""
        self.load_data()
        current_page = self.current_page
        if current_page == "view":
            self.show_view_page()
        elif current_page == "dashboard":
            self.update_dashboard_page()
        elif current_page == "reminders":
            self.refresh_reminder_log()
        else:
            self.refresh_appointments()
        self.show_toast("All data refreshed!", "success")

    def save_data(self):
        """Save appointments data to file"""
        try:
            with open(APPOINTMENTS_FILE, 'w') as f:
                json.dump(self.appointments, f, indent=4, default=str)
        except Exception as e:
            self.show_toast(f"Auto-save failed: {e}", "error")

    def load_data(self):
        """Load appointments data from file - FIXED VERSION"""
        try:
            if os.path.exists(APPOINTMENTS_FILE):
                with open(APPOINTMENTS_FILE, 'r') as f:
                    raw_appointments = json.load(f)
                    # Validate and clean each appointment record
                    self.appointments = []
                    for apt in raw_appointments:
                        if isinstance(apt, dict):  # Ensure it's a dictionary
                            cleaned_apt = self.validate_appointment_data(apt)
                            self.appointments.append(cleaned_apt)
            self.update_stats()
        except (json.JSONDecodeError, FileNotFoundError) as e:
            self.appointments = []
            self.show_toast(f"Could not load appointments: {e}", "warning")

    def create_reminders_page(self):
        """Create auto-reminders management page"""
        theme = self.get_theme()
        page = tk.Frame(self.content_area, bg=theme["bg_primary"])
        self.pages["reminders"] = page
        header = tk.Frame(page, bg=theme["bg_primary"])
        header.pack(fill='x', pady=(0, 20))
        tk.Label(header, text="📱 Auto Reminder System", font=self.fonts["title"], bg=theme["bg_primary"], fg=theme["text_primary"]).pack(side='left')
        controls_frame = tk.Frame(header, bg=theme["bg_primary"])
        controls_frame.pack(side='right')
        tk.Button(controls_frame, text="🔔 ON" if self.reminder_settings["enabled"] else "🔕 OFF", command=self.toggle_reminder_system, bg=theme["success"] if self.reminder_settings["enabled"] else theme["danger"], fg="white").pack(side='left', padx=5)
        tk.Button(controls_frame, text="📱 AUTO" if self.reminder_settings.get("auto_send_whatsapp", True) else "📱 MANUAL", command=self.toggle_whatsapp_auto_send, bg=theme["accent"] if self.reminder_settings.get("auto_send_whatsapp", True) else theme["warning"], fg="white").pack(side='left', padx=5)
        settings_frame = tk.LabelFrame(page, text="⚙️ Reminder Settings", bg=theme["bg_secondary"], fg=theme["text_primary"], font=self.fonts["body"])
        settings_frame.pack(fill='x', padx=5, pady=5)
        settings_inner = tk.Frame(settings_frame, bg=theme["bg_secondary"])
        settings_inner.pack(fill='x', padx=15, pady=15)
        self.reminder_vars = {}
        options = [("remind_3_days", "📅 3 Days Before"), ("remind_1_day", "📅 1 Day Before"), ("remind_morning", "🌅 Morning Of"), ("remind_1_hour", "⏰ 1 Hour Before")]
        for i, (key, text) in enumerate(options):
            var = tk.BooleanVar(value=self.reminder_settings.get(key, True))
            self.reminder_vars[key] = var
            tk.Checkbutton(settings_inner, text=text, variable=var, bg=theme["bg_secondary"], fg=theme["text_primary"], font=self.fonts["body"], anchor='w', command=self.save_reminder_settings).grid(row=i, column=0, sticky='w', pady=2)

        # Log section
        log_frame = tk.LabelFrame(page, text="📜 Activity Log", bg=theme["bg_secondary"], fg=theme["text_primary"], font=self.fonts["body"])
        log_frame.pack(fill='both', expand=True, padx=5, pady=5)
        log_scroll = ttk.Scrollbar(log_frame)
        log_scroll.pack(side='right', fill='y')
        self.reminder_log_tree = ttk.Treeview(log_frame, columns=('Time', 'Patient', 'Contact', 'Activity', 'Status'), show='headings', yscrollcommand=log_scroll.set, height=8)
        log_scroll.config(command=self.reminder_log_tree.yview)
        for col, width in [('Time', 100), ('Patient', 120), ('Contact', 120), ('Activity', 150), ('Status', 80)]:
            self.reminder_log_tree.heading(col, text=col)
            self.reminder_log_tree.column(col, width=width, anchor='center')
        self.reminder_log_tree.pack(fill='both', expand=True)

    def start_reminder_system(self):
        if not self.reminder_running and self.reminder_settings["enabled"]:
            self.reminder_running = True
            self.reminder_thread = threading.Thread(target=self.reminder_loop, daemon=True)
            self.reminder_thread.start()
            self.log_reminder_activity("System", "", "Reminder system started", "SUCCESS")

    def stop_reminder_system(self):
        self.reminder_running = False
        if self.reminder_thread:
            self.reminder_thread = None
        self.log_reminder_activity("System", "", "Reminder system stopped", "INFO")

    def reminder_loop(self):
        while self.reminder_running:
            try:
                if self.reminder_settings["enabled"]:
                    self.check_and_send_reminders()
                time.sleep(self.reminder_settings.get("check_interval", 300))
            except Exception as e:
                self.log_reminder_activity("System", "", f"Error in loop: {str(e)}", "ERROR")
                time.sleep(60)

    def check_and_send_reminders(self):
        now = datetime.now()
        if not self.is_business_hours(now.strftime("%H:%M")): 
            return
        for appointment in self.appointments:
            if not appointment.get('enable_reminders', True) and not appointment.get('enable_email', True): 
                continue
            apt_datetime = self.get_appointment_datetime(appointment)
            if not apt_datetime: 
                continue
            time_diff = apt_datetime - now
            self.check_reminder_type(appointment, time_diff, "3_days", timedelta(days=3))
            self.check_reminder_type(appointment, time_diff, "1_day", timedelta(days=1))
            self.check_reminder_type(appointment, time_diff, "morning", timedelta(hours=12))
            self.check_reminder_type(appointment, time_diff, "1_hour", timedelta(hours=1))

    def check_reminder_type(self, appointment, time_diff, reminder_type, target_time):
        setting_key = f"remind_{reminder_type}"
        if not self.reminder_settings.get(setting_key, True): 
            return
        reminder_key = f"{appointment.get('id', 'unknown')}_{reminder_type}"
        if reminder_key in self.sent_reminders: 
            return
        should_send = False
        if reminder_type == "morning" and time_diff.days == 0 and 8 <= datetime.now().hour < 10: 
            should_send = True
        elif reminder_type == "1_hour" and timedelta(minutes=30) <= time_diff <= timedelta(minutes=90): 
            should_send = True
        elif reminder_type in ["3_days", "1_day"] and abs(time_diff - target_time) <= timedelta(hours=6): 
            should_send = True
        if should_send:
            whatsapp_success = self.send_auto_whatsapp_reminder(appointment, reminder_type) if appointment.get('enable_reminders') else False
            email_success = self.send_email_reminder(appointment, reminder_type) if appointment.get('enable_email') else False
            if whatsapp_success or email_success:
                self.sent_reminders[reminder_key] = datetime.now().isoformat()
                self.save_reminder_data()

    def send_auto_whatsapp_reminder(self, appointment, reminder_type):
        try:
            message = self.get_reminder_message(appointment, reminder_type)
            phone = appointment.get('phone_number', '')
            clean_phone = self.clean_phone_number(phone)
            if not clean_phone:
                self.log_reminder_activity(appointment.get('patient_name', 'Unknown'), phone, "Invalid phone number", "ERROR")
                return False
            if self.reminder_settings.get("auto_send_whatsapp", True):
                if self.open_whatsapp_chat(clean_phone, message, auto_send=True):
                    self.log_reminder_activity(appointment.get('patient_name', 'Unknown'), clean_phone, f"{reminder_type} reminder sent", "SENT ✅")
                    self.root.after(0, lambda: self.show_whatsapp_notification(appointment, reminder_type))
                    time.sleep(self.reminder_settings.get("whatsapp_delay", 3))
                    return True
                else:
                    self.log_reminder_activity(appointment.get('patient_name', 'Unknown'), clean_phone, f"Failed to send {reminder_type}", "FAILED ❌")
                    return False
            else:
                self.log_reminder_activity(appointment.get('patient_name', 'Unknown'), clean_phone, f"{reminder_type} (auto-send off)", "LOGGED 📝")
                return True
        except Exception as e:
            self.log_reminder_activity(appointment.get('patient_name', 'Unknown'), appointment.get('phone_number', 'N/A'), f"Error sending {reminder_type}: {e}", "ERROR ❌")
            return False

    def open_whatsapp_chat(self, phone_number, message, auto_send=False):
        """
        Opens WhatsApp chat. Can attempt to open the desktop app or web.
        auto_send only works with pyautogui (not ideal for all systems).
        """
        encoded_message = urllib.parse.quote(message)
        
        if self.reminder_settings.get("whatsapp_app_mode", False):
            # Try to open WhatsApp desktop app
            try:
                # This URI scheme is often registered by WhatsApp desktop app
                if platform.system() == "Windows":
                    # On Windows, try to use start command for whatsapp:// protocol
                    subprocess.Popen(f'start whatsapp://send?phone={phone_number}&text={encoded_message}', shell=True)
                elif platform.system() == "Darwin":
                    # On macOS, use open command
                    subprocess.Popen(['open', f'whatsapp://send?phone={phone_number}&text={encoded_message}'])
                else:
                    # Generic fallback for Linux, might require xdg-open or similar setup
                    webbrowser.open(f"whatsapp://send?phone={phone_number}&text={encoded_message}")

                if auto_send:
                    # Give time for the app to open and then simulate enter key
                    self.root.after(5000, self.auto_send_whatsapp_message) # Adjust delay as needed
                return True
            except Exception as e:
                print(f"Failed to open WhatsApp desktop app: {e}. Falling back to web.")
                webbrowser.open(f"https://wa.me/{phone_number}?text={encoded_message}")
                if auto_send:
                    self.root.after(5000, self.auto_send_whatsapp_message)
                return True # Consider it successful if web opens
        else:
            # Default to WhatsApp Web
            webbrowser.open(f"https://web.whatsapp.com/send?phone={phone_number}&text={encoded_message}")
            if auto_send:
                self.root.after(8000, self.auto_send_whatsapp_message) # Give whatsapp web time to load
            return True
        
        return False # Should not be reached

    def auto_send_whatsapp_message(self):
        """Simulate Enter key press to send WhatsApp message. Requires 'pyautogui'."""
        try:
            import pyautogui
            # A longer delay is often needed for WhatsApp Web to fully load and focus the chat input
            # For desktop app, delay might be shorter but still needed for focus
            time.sleep(5) 
            pyautogui.press('enter')
            self.show_toast("WhatsApp message auto-sent (requires focus on chat).", "info")
        except ImportError:
            self.show_toast("PyAutoGUI not installed. Cannot auto-send message.", "warning")
            print("PyAutoGUI not installed. Cannot auto-send message.")
        except Exception as e:
            self.show_toast(f"Failed to auto-press Enter: {e}", "error")
            print(f"Failed to auto-press Enter: {e}")

    def clean_phone_number(self, phone):
        if not phone or not phone.strip(): 
            return None
        clean = re.sub(r'[^\d+]', '', phone.strip())
        # If it doesn't start with '+', add a default country code.
        # This assumes a default, which might not be correct for all users.
        # Consider making this configurable.
        if not clean.startswith('+'):
            clean = '+' + clean.lstrip('0') 
        return clean.replace('+', '') # Remove the '+' for the wa.me link which expects digits only

    def get_reminder_message(self, appointment, reminder_type):
        name = appointment.get('patient_name', 'Patient')
        procedure = appointment.get('procedure', 'N/A')
        apt_date = appointment.get('appointment_date', 'N/A')
        apt_time = appointment.get('appointment_time', '09:00')
        messages = {
            "3_days": f"🏥 Hi {name}! Reminder for your {procedure} appointment in 3 days on {apt_date} at {apt_time}. Please confirm. Thank you! 😊",
            "1_day": f"🏥 Hello {name}! Your {procedure} appointment is tomorrow, {apt_date} at {apt_time}. See you soon! 👋",
            "morning": f"🌅 Good morning {name}! You have a {procedure} appointment TODAY at {apt_time}. See you soon!",
            "1_hour": f"⏰ Hi {name}! Your {procedure} appointment is in 1 HOUR at {apt_time}. Please make your way to our clinic."
        }
        return messages.get(reminder_type, f"Hi {name}, reminder about your {procedure} appointment.")

    def show_whatsapp_notification(self, appointment, reminder_type):
        self.show_toast(f"📱 Reminder sent to {appointment.get('patient_name', 'Patient')}", "success")

    def toggle_whatsapp_auto_send(self):
        self.reminder_settings["auto_send_whatsapp"] = not self.reminder_settings.get("auto_send_whatsapp", True)
        self.save_reminder_settings()
        self.show_toast(f"Auto WhatsApp sending {'enabled' if self.reminder_settings['auto_send_whatsapp'] else 'disabled'}!", "info")
        self.update_reminder_status()

    def get_appointment_datetime(self, appointment):
        try:
            return datetime.combine(
                datetime.strptime(appointment.get('appointment_date', ''), '%Y-%m-%d').date(), 
                datetime.strptime(appointment.get('appointment_time', '09:00'), '%H:%M').time()
            )
        except (ValueError, KeyError):
            return None

    def is_business_hours(self, current_time_str):
        try:
            start_time = datetime.strptime(self.reminder_settings["business_hours_start"], '%H:%M').time()
            end_time = datetime.strptime(self.reminder_settings["business_hours_end"], '%H:%M').time()
            current_time = datetime.strptime(current_time_str, '%H:%M').time()
            return start_time <= current_time <= end_time
        except (ValueError, KeyError):
            return True

    def toggle_reminder_system(self):
        self.reminder_settings["enabled"] = not self.reminder_settings["enabled"]
        if self.reminder_settings["enabled"]:
            self.start_reminder_system()
            self.show_toast("Auto-reminder system enabled! 🔔", "success")
        else:
            self.stop_reminder_system()
            self.show_toast("Auto-reminder system disabled! 🔕", "warning")
        self.save_reminder_settings()
        self.update_reminder_status()

    def update_reminder_status(self):
        theme = self.get_theme()
        if hasattr(self, 'reminder_status'):
            self.reminder_status.config(text=f"🔔 Reminders: {'ON' if self.reminder_settings['enabled'] else 'OFF'}", bg=theme["success"] if self.reminder_settings["enabled"] else theme["danger"])
        if hasattr(self, 'whatsapp_status'):
            self.whatsapp_status.config(text="📱 WhatsApp AUTO" if self.reminder_settings.get("auto_send_whatsapp", True) else "📱 Manual Only", bg=theme["accent"] if self.reminder_settings.get("auto_send_whatsapp", True) else theme["warning"])

    def log_reminder_activity(self, patient, contact, activity, status):
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        try:
            with open(REMINDER_LOG_FILE, 'a', encoding='utf-8') as f:
                f.write(f"{timestamp} | {patient} | {contact} | {activity} | {status}\n")
        except Exception as e:
            print(f"Error logging activity: {e}")
        if hasattr(self, 'reminder_log_tree'):
            self.root.after(0, lambda: self.add_log_entry(timestamp, patient, contact, activity, status))

    def add_log_entry(self, timestamp, patient, contact, activity, status):
        try:
            self.reminder_log_tree.insert('', 0, values=(timestamp.split(' ')[1], patient[:15], contact[:15], activity[:20], status))
            if len(self.reminder_log_tree.get_children()) > 100:
                self.reminder_log_tree.delete(self.reminder_log_tree.get_children()[-1])
        except Exception:
            pass

    def refresh_reminder_log(self):
        if not hasattr(self, 'reminder_log_tree'):
            return
        for item in self.reminder_log_tree.get_children(): 
            self.reminder_log_tree.delete(item)
        try:
            if os.path.exists(REMINDER_LOG_FILE):
                with open(REMINDER_LOG_FILE, 'r', encoding='utf-8') as f:
                    for line in reversed(f.readlines()[-100:]):
                        parts = line.strip().split(' | ')
                        if len(parts) >= 5: 
                            self.add_log_entry(parts[0], parts[1], parts[2], parts[3], parts[4])
        except Exception as e:
            print(f"Error refreshing log: {e}")

    def save_reminder_settings(self):
        """Saves reminder settings from the UI to the file."""
        try:
            # Update the main settings dict from the UI variables
            if hasattr(self, 'reminder_settings_vars'):
                for key, var in self.reminder_settings_vars.items():
                    self.reminder_settings[key] = var.get()
            
            with open(REMINDER_SETTINGS_FILE, 'w') as f:
                 json.dump(self.reminder_settings, f, indent=4)
            self.show_toast("Reminder settings saved!", "success")
            self.update_reminder_status() # Reflect changes in sidebar
        except Exception as e:
            self.show_toast(f"Error saving reminder settings: {e}", "error")

    def load_reminder_data(self):
        try:
            if os.path.exists(REMINDER_SETTINGS_FILE):
                with open(REMINDER_SETTINGS_FILE, 'r') as f: 
                    self.reminder_settings.update(json.load(f))
        except Exception: 
            pass
        try:
            if os.path.exists(SENT_REMINDERS_FILE):
                with open(SENT_REMINDERS_FILE, 'r') as f: 
                    self.sent_reminders = json.load(f)
        except Exception: 
            pass

    def save_reminder_data(self):
        try:
            with open(SENT_REMINDERS_FILE, 'w') as f: 
                json.dump(self.sent_reminders, f, indent=4)
        except Exception as e:
            self.show_toast(f"Error saving reminder data: {e}", "error")

    def create_stat_card(self, parent, title, icon, var_name):
        theme = self.get_theme()
        card = tk.Frame(parent, bg=theme["accent"], relief='solid', bd=1, highlightbackground=theme["bg_accent"], highlightthickness=1)
        card.pack(side='left', fill='x', expand=True, padx=5, pady=5)
        
        icon_label = tk.Label(card, text=icon, font=("Arial", 24), bg=theme["accent"], fg="white")
        icon_label.pack(side='left', padx=(15, 10), pady=10)
        
        text_frame = tk.Frame(card, bg=theme["accent"])
        text_frame.pack(side='left', fill='x', expand=True, pady=10, padx=(0, 15))
        
        value_label = tk.Label(text_frame, text="0", font=self.fonts["title"], bg=theme["accent"], fg="white", anchor='w')
        value_label.pack(fill='x')
        
        title_label = tk.Label(text_frame, text=title, font=self.fonts["small"], bg=theme["accent"], fg="#e0e0e0", anchor='w')
        title_label.pack(fill='x')

        setattr(self, f"{var_name}_label", value_label)

    def create_settings_page(self):
        """Creates a comprehensive settings page with multiple sections."""
        theme = self.get_theme()
        page = tk.Frame(self.content_area, bg=theme["bg_primary"])
        self.pages["settings"] = page

        tk.Label(page, text="⚙️ Application Settings", font=self.fonts["title"], bg=theme["bg_primary"], fg=theme["text_primary"]).pack(pady=(0, 10), anchor='w', padx=10)

        # Create a canvas and a scrollbar
        canvas = tk.Canvas(page, bg=theme["bg_primary"], highlightthickness=0)
        scrollbar = ttk.Scrollbar(page, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=theme["bg_primary"])

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # --- START: MOUSE WHEEL SCROLL BINDING ---
        def _on_mousewheel(event):
            """Cross-platform mouse wheel scroll handler."""
            if platform.system() == "Windows":
                canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
            elif platform.system() == "Darwin":  # macOS
                canvas.yview_scroll(int(-1 * event.delta), "units")
            else:  # Linux and other systems
                if event.num == 4:
                    canvas.yview_scroll(-1, "units")
                elif event.num == 5:
                    canvas.yview_scroll(1, "units")

        def _bind_scroll(event):
            """Bind mouse wheel events for scrolling."""
            canvas.bind_all("<MouseWheel>", _on_mousewheel)
            canvas.bind_all("<Button-4>", _on_mousewheel)
            canvas.bind_all("<Button-5>", _on_mousewheel)

        def _unbind_scroll(event):
            """Unbind mouse wheel events."""
            canvas.unbind_all("<MouseWheel>")
            canvas.unbind_all("<Button-4>")
            canvas.unbind_all("<Button-5>")

        # Bind events when mouse enters the canvas, unbind when it leaves.
        # This ensures this specific canvas scrolls only when the mouse is over it.
        canvas.bind('<Enter>', _bind_scroll)
        canvas.bind('<Leave>', _unbind_scroll)
        # --- END: MOUSE WHEEL SCROLL BINDING ---

        canvas.pack(side="left", fill="both", expand=True, padx=5)
        scrollbar.pack(side="right", fill="y")

        # --- Email Settings Section ---
        email_frame = ttk.LabelFrame(scrollable_frame, text="📧 Email & Clinic Settings", padding=(10, 10))
        email_frame.pack(fill='x', padx=10, pady=10)
        
        self.email_settings_vars = {}
        email_fields = [
            ("SMTP Server:", "smtp_server"), ("SMTP Port:", "smtp_port"),
            ("Your Email:", "email_address"), ("App Password:", "app_password"),
            ("Clinic Name:", "clinic_name"), ("Clinic Address:", "clinic_address"),
            ("Clinic Phone:", "clinic_phone")
        ]

        for i, (label, key) in enumerate(email_fields):
            tk.Label(email_frame, text=label, font=self.fonts["body"]).grid(row=i, column=0, sticky='w', pady=2)
            var = tk.StringVar()
            self.email_settings_vars[key] = var
            show_char = "*" if "password" in key else None
            entry = tk.Entry(email_frame, textvariable=var, font=self.fonts["body"], width=40, show=show_char)
            entry.grid(row=i, column=1, sticky='ew', pady=2, padx=5)

        self.email_settings_vars["auto_send_email"] = tk.BooleanVar()
        tk.Checkbutton(email_frame, text="Automatically send email reminders", variable=self.email_settings_vars["auto_send_email"], font=self.fonts["body"]).grid(row=len(email_fields), column=0, columnspan=2, sticky='w', pady=5)
        
        email_btn_frame = tk.Frame(email_frame)
        email_btn_frame.grid(row=len(email_fields) + 1, column=0, columnspan=2, pady=10)
        tk.Button(email_btn_frame, text="💾 Save Email Settings", command=self.save_email_settings, bg=theme["success"], fg='white', relief='flat').pack(side='left', padx=5)
        tk.Button(email_btn_frame, text="🧪 Test Connection", command=self.test_email_connection, bg=theme["accent"], fg='white', relief='flat').pack(side='left', padx=5)

        # --- WhatsApp & Automation Settings Section ---
        reminder_frame = ttk.LabelFrame(scrollable_frame, text="📱 WhatsApp & Automation Settings", padding=(10, 10))
        reminder_frame.pack(fill='x', padx=10, pady=10)
        
        self.reminder_settings_vars = {}
        
        # Initialize the BooleanVar *before* using it in the Checkbutton
        auto_whatsapp_var = tk.BooleanVar()
        self.reminder_settings_vars["auto_send_whatsapp"] = auto_whatsapp_var
        tk.Checkbutton(reminder_frame, text="Automatically send WhatsApp messages", variable=auto_whatsapp_var, font=self.fonts["body"]).grid(row=0, column=0, columnspan=2, sticky='w', pady=5)

        # New checkbox for WhatsApp app mode
        whatsapp_app_mode_var = tk.BooleanVar()
        self.reminder_settings_vars["whatsapp_app_mode"] = whatsapp_app_mode_var
        tk.Checkbutton(reminder_frame, text="Attempt to open WhatsApp Desktop App (instead of Web)", variable=whatsapp_app_mode_var, font=self.fonts["body"]).grid(row=1, column=0, columnspan=2, sticky='w', pady=5)


        reminder_fields = [
            ("Business Hours Start (HH:MM):", "business_hours_start"),
            ("Business Hours End (HH:MM):", "business_hours_end"),
            ("Check Interval (seconds):", "check_interval"),
            ("WhatsApp Send Delay (seconds):", "whatsapp_delay")
        ]
        for i, (label, key) in enumerate(reminder_fields, 2): # Adjusted row start due to new checkbox
            tk.Label(reminder_frame, text=label, font=self.fonts["body"]).grid(row=i, column=0, sticky='w', pady=2)
            var = tk.StringVar()
            self.reminder_settings_vars[key] = var
            tk.Entry(reminder_frame, textvariable=var, font=self.fonts["body"], width=10).grid(row=i, column=1, sticky='w', pady=2, padx=5)

        tk.Button(reminder_frame, text="💾 Save Automation Settings", command=self.save_reminder_settings, bg=theme["success"], fg='white', relief='flat').grid(row=len(reminder_fields)+2, column=0, columnspan=2, pady=10)
        
        # --- Display Settings ---
        display_frame = ttk.LabelFrame(scrollable_frame, text="🎨 Display Settings", padding=(10, 10))
        display_frame.pack(fill='x', padx=10, pady=10)
        
        tk.Label(display_frame, text="Theme:", font=self.fonts["body"]).pack(side='left', padx=(0,10))
        tk.Button(display_frame, text="☀️ Light Mode", command=lambda: self.switch_theme("light"), bg=theme["bg_accent"], fg=theme["text_primary"], relief='flat').pack(side='left', padx=5)
        tk.Button(display_frame, text="🌙 Dark Mode", command=lambda: self.switch_theme("dark"), bg=theme["bg_accent"], fg=theme["text_primary"], relief='flat').pack(side='left', padx=5)

        # --- Data Management Section ---
        data_frame = ttk.LabelFrame(scrollable_frame, text="💾 Data Management", padding=(10, 10))
        data_frame.pack(fill='x', padx=10, pady=10, expand=True)
        
        export_btns = [
            ("Export to CSV", self.export_to_csv),
            ("Export to Excel", self.export_to_excel),
            ("Export to JSON", self.export_data),
            ("Import from JSON", self.import_data)
        ]
        for i, (text, command) in enumerate(export_btns):
            if "Excel" in text and not PANDAS_AVAILABLE: 
                continue
            tk.Button(data_frame, text=text, command=command, bg=theme["accent"], fg="white", relief='flat').pack(side='left', padx=5, pady=5)

    def update_settings_page_vars(self):
        """Update the variables on the settings page with loaded values."""
        if hasattr(self, 'email_settings_vars'):
            for key, var in self.email_settings_vars.items():
                var.set(self.email_settings.get(key, ""))
        
        if hasattr(self, 'reminder_settings_vars'):
            for key, var in self.reminder_settings_vars.items():
                # For boolean vars, ensure we set them correctly
                if isinstance(var, tk.BooleanVar):
                    var.set(self.reminder_settings.get(key, False))
                else:
                    var.set(self.reminder_settings.get(key, ""))

    def test_email_connection(self):
        """Test email connection using settings from the UI with enhanced SSL support."""
        try:
            if not EMAIL_AVAILABLE:
                self.show_toast("Please install yagmail: pip install yagmail", "error")
                return

            # Get values directly from the UI entry fields
            email_addr = self.email_settings_vars["email_address"].get()
            app_password = self.email_settings_vars["app_password"].get()
            smtp_server = self.email_settings_vars["smtp_server"].get()
            smtp_port = self.email_settings_vars["smtp_port"].get()

            if not all([email_addr, app_password, smtp_server, smtp_port]):
                self.show_toast("Please fill all email fields before testing", "warning")
                return
            
            # Try different SSL/TLS configurations for private servers
            port = int(smtp_port)
            
            # Method 1: Try with explicit SSL configuration
            try:
                import smtplib
                import ssl
                
                # Create SSL context with more permissive settings for private servers
                context = ssl.create_default_context()
                context.check_hostname = False
                context.verify_mode = ssl.CERT_NONE
                
                # Try different connection methods based on port
                server = None
                
                if port == 465:  # SSL
                    server = smtplib.SMTP_SSL(smtp_server, port, context=context)
                elif port == 587:  # TLS
                    server = smtplib.SMTP(smtp_server, port)
                    server.starttls(context=context)
                else:  # Try both
                    try:
                        # First try as SSL
                        server = smtplib.SMTP_SSL(smtp_server, port, context=context)
                    except:
                        # Then try as TLS
                        server = smtplib.SMTP(smtp_server, port)
                        server.starttls(context=context)
                
                # Login and send test email
                server.login(email_addr, app_password)
                
                # Compose test message
                from email.mime.text import MIMEText
                from email.mime.multipart import MIMEMultipart
                
                msg = MIMEMultipart()
                msg['From'] = email_addr
                msg['To'] = email_addr
                msg['Subject'] = "Test Email - Clinic System"
                
                body = "This is a test email from your Clinic System. Connection is successful!"
                msg.attach(MIMEText(body, 'plain'))
                
                server.send_message(msg)
                server.quit()
                
                self.show_toast("Test email sent successfully using direct SMTP!", "success")
                return
                
            except Exception as smtp_error:
                print(f"Direct SMTP failed: {smtp_error}")
                
                # Method 2: Try with yagmail but different SSL settings
                try:
                    # Configure yagmail for private servers
                    yag = yagmail.SMTP(
                        user=email_addr,
                        password=app_password,
                        host=smtp_server,
                        port=port,
                        smtp_starttls=True if port == 587 else False,
                        smtp_ssl=True if port == 465 else False,
                        smtp_set_debuglevel=0,
                        smtp_skip_login=False
                    )
                    
                    yag.send(
                        to=email_addr, 
                        subject="Test Email - Clinic System", 
                        contents="This is a test email from your Clinic System. Connection is successful!"
                    )
                    yag.close()
                    self.show_toast("Test email sent successfully using yagmail!", "success")
                    return
                    
                except Exception as yag_error:
                    print(f"Yagmail method failed: {yag_error}")
                    
                    # Method 3: Try without SSL/TLS (for local/internal servers)
                    try:
                        server = smtplib.SMTP(smtp_server, port)
                        server.login(email_addr, app_password)
                        
                        msg = MIMEMultipart()
                        msg['From'] = email_addr
                        msg['To'] = email_addr
                        msg['Subject'] = "Test Email - Clinic System"
                        body = "This is a test email from your Clinic System. Connection is successful!"
                        msg.attach(MIMEText(body, 'plain'))
                        
                        server.send_message(msg)
                        server.quit()
                        
                        self.show_toast("Test email sent successfully without SSL!", "success")
                        return
                        
                    except Exception as no_ssl_error:
                        print(f"No SSL method failed: {no_ssl_error}")
                        
                        # If all methods fail, show detailed error
                        error_details = f"""
    Email test failed with all methods:

    1. SMTP SSL/TLS Error: {smtp_error}
    2. Yagmail Error: {yag_error}  
    3. No SSL Error: {no_ssl_error}

    For private email servers, try:
    - Port 25 (no encryption)
    - Port 587 (STARTTLS)
    - Port 465 (SSL)
    - Check if your server requires specific auth methods
                        """
                        self.show_toast("All email methods failed. Check console for details.", "error", duration=8000)
                        print(error_details)
                        
        except Exception as e:
            self.show_toast(f"Email configuration error: {str(e)}", "error", duration=8000)
            print(f"Email test error: {e}")

    def create_notification_system(self):
        """Placeholder for notification system initialization."""
        pass

    def switch_theme(self, theme_name):
        """Switches the theme only if it's different from the current one."""
        if self.current_theme != theme_name:
            self.toggle_theme()

    def toggle_theme(self):
        """Toggle between light and dark themes."""
        self.current_theme = "dark" if self.current_theme == "light" else "light"
        # Recreate the UI to apply the theme thoroughly
        for widget in self.root.winfo_children():
            widget.destroy()
        self.setup_main_layout()
        self.create_sidebar()
        self.create_main_content()
        self.load_email_settings()
        self.load_reminder_data()
        self.update_settings_page_vars()
        self.show_toast(f"Switched to {self.current_theme} theme", "success")
        # Restore the last active page
        page_command = getattr(self, f"show_{self.current_page}_page")
        self.navigate_to(page_command, self.current_page)

    def setup_keyboard_shortcuts(self):
        self.root.bind('<Control-s>', lambda e: self.save_data())
        self.root.bind('<Control-n>', lambda e: self.navigate_to(self.show_add_page, "add"))
        self.root.bind('<Control-r>', lambda e: self.refresh_all())
        self.root.bind('<F1>', lambda e: self.show_help())

    def start_auto_save(self):
        if self.auto_save_active:
            self.save_data()
            self.root.after(30000, self.start_auto_save)

    def send_manual_whatsapp(self):
        selected_items = self.appointments_tree.selection()
        if not selected_items:
            self.show_toast("Please select at least one appointment to send WhatsApp to!", "warning")
            return

        for selected_id in selected_items:
            item = self.appointments_tree.item(selected_id)
            apt_id = item['values'][0]
            appointment = next((apt for apt in self.appointments if str(apt.get('id', '')) == str(apt_id)), None)

            if appointment:
                message = f"Hi {appointment.get('patient_name', 'Patient')}, this is a manual reminder for your appointment on {appointment.get('appointment_date', 'N/A')} at {appointment.get('appointment_time', '09:00')}. Please confirm if you will attend. Thank you!"
                phone = self.clean_phone_number(appointment.get('phone_number', ''))

                if phone:
                    if self.open_whatsapp_chat(phone, message): # Use the centralized function
                        self.log_reminder_activity(appointment.get('patient_name', 'Unknown'), phone, "Manual WhatsApp sent", "MANUAL ✅")
                    else:
                        self.show_toast(f"Failed to open WhatsApp for {appointment.get('patient_name')}", "error")
                else:
                    self.show_toast(f"Invalid phone number for {appointment.get('patient_name')}", "warning")
            else:
                self.show_toast("Appointment not found for selected item.", "error")


    def send_manual_email(self):
        """Send manual email to selected patient(s)."""
        selected_items = self.appointments_tree.selection()
        if not selected_items:
            self.show_toast("Please select at least one appointment to send email to!", "warning")
            return

        for selected_id in selected_items:
            item = self.appointments_tree.item(selected_id)
            apt_id = item['values'][0]
            appointment = next((apt for apt in self.appointments if str(apt.get('id', '')) == str(apt_id)), None)

            if appointment:
                email = appointment.get('email', '').strip()
                if email and self.validate_email(email):
                    # We can use the existing send_email_reminder but provide a 'manual' type
                    if self.send_email_reminder(appointment, "manual"):
                        self.show_toast(f"Email sent to {appointment.get('patient_name')}", "success")
                    else:
                        self.show_toast(f"Failed to send email to {appointment.get('patient_name')}", "error")
                else:
                    self.show_toast(f"Invalid or missing email for {appointment.get('patient_name')}", "warning")
            else:
                self.show_toast("Appointment not found for selected item.", "error")

    def edit_appointment(self):
        selected = self.appointments_tree.selection()
        if not selected:
            self.show_toast("Please select an appointment to edit!", "warning")
            return
        item = self.appointments_tree.item(selected[0])
        apt_id = item['values'][0]
        appointment = next((apt for apt in self.appointments if str(apt.get('id', '')) == str(apt_id)), None)
        if appointment: 
            self.create_edit_dialog(appointment)

    def create_edit_dialog(self, appointment):
        """Create edit dialog, showing a read-only Patient ID and editable Email."""
        theme = self.get_theme()
        dialog = tk.Toplevel(self.root)
        dialog.title("✏️ Edit Appointment")
        dialog.geometry("450x650") # Increased height to accommodate new fields
        dialog.configure(bg=theme["bg_primary"])
        dialog.grab_set()
        dialog.resizable(False, False)
        dialog.transient(self.root)

        header = tk.Frame(dialog, bg=theme["accent"], height=60)
        header.pack(fill='x')
        header.pack_propagate(False)
        tk.Label(header, text="✏️ Edit Patient Appointment", font=self.fonts["heading"], bg=theme["accent"], fg="white").pack(expand=True)

        form_frame = tk.Frame(dialog, bg=theme["bg_primary"])
        form_frame.pack(fill='both', expand=True, padx=20, pady=20)
        edit_vars = {}
        row = 0

        tk.Label(form_frame, text="🆔 Patient ID:", font=self.fonts["body"], bg=theme["bg_primary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='w', pady=5)
        id_entry = tk.Entry(form_frame, font=self.fonts["body"], bg=theme["bg_accent"], fg=theme["text_secondary"], relief='flat', bd=1)
        id_entry.insert(0, appointment.get('id', 'N/A'))
        id_entry.config(state='readonly')
        id_entry.grid(row=row, column=1, sticky='ew', pady=5, padx=(10, 0))
        row += 1

        # Parse procedure to separate modality and details for the edit dialog
        modality, procedure_details = self.parse_procedure(appointment.get('procedure', ''))

        # Modified: Update fields list to separate modality and procedure
        fields = [
            ("👤 Patient Name:", "patient_name", appointment.get('patient_name', '')),
            ("🔬 Modality:", "modality", modality),  # Changed to store just the modality
            ("📝 Procedure Details:", "procedure_details", procedure_details),  # New field for details
            ("📱 WhatsApp Phone:", "phone_number", appointment.get('phone_number', '')),
            ("📧 Email:", "email", appointment.get('email', '')),
            ("💳 Health Card:", "health_card", appointment.get('health_card', '')),
            ("📅 Appointment Date:", "appointment_date", appointment.get('appointment_date', '')),
            ("⏰ Appointment Time:", "appointment_time", appointment.get('appointment_time', '09:00')),
            ("🏥 Clinic Appt. Date:", "clinic_appointment_date", appointment.get('clinic_appointment_date', ''))
        ]
        
        for i, (label, key, value) in enumerate(fields, start=row):
            tk.Label(form_frame, text=label, font=self.fonts["body"], bg=theme["bg_primary"], fg=theme["text_primary"], anchor='w').grid(row=i, column=0, sticky='w', pady=5)
            
            if key == "health_card":
                var = tk.StringVar(value=value if value else HEALTH_CARD_OPTIONS[0])
                edit_vars[key] = var
                ttk.Combobox(form_frame, textvariable=var, values=HEALTH_CARD_OPTIONS, state="readonly", font=self.fonts["body"]).grid(row=i, column=1, sticky='ew', pady=5, padx=(10, 0))
            elif key == "modality":
                # Create a dropdown for modality selection
                var = tk.StringVar(value=value if value else "")
                edit_vars[key] = var
                ttk.Combobox(form_frame, textvariable=var, values=[p for p in PROCEDURE_TYPES if p != "All"], state="readonly", font=self.fonts["body"]).grid(row=i, column=1, sticky='ew', pady=5, padx=(10, 0))
            elif key in ["appointment_date", "clinic_appointment_date"]:
                date_entry_widget = DateEntry(form_frame, width=15, background=theme["accent"], foreground='white', borderwidth=1, font=self.fonts["body"], mindate=date.today())
                try:
                    if value:
                        current_date = datetime.strptime(value, '%Y-%m-%d').date()
                        date_entry_widget.set_date(current_date)
                    else:
                        date_entry_widget.set_date(None)
                except ValueError:
                    date_entry_widget.set_date(None)
                date_entry_widget.grid(row=i, column=1, sticky='ew', pady=5, padx=(10, 0))
                edit_vars[key] = date_entry_widget
            else:
                var = tk.StringVar(value=value)
                edit_vars[key] = var
                tk.Entry(form_frame, textvariable=var, font=self.fonts["body"], bg=theme["bg_secondary"], fg=theme["text_primary"], relief='flat', bd=1).grid(row=i, column=1, sticky='ew', pady=5, padx=(10, 0))
        row += len(fields)

        # Notes field (Text widget)
        tk.Label(form_frame, text="📝 Notes:", font=self.fonts["body"], bg=theme["bg_primary"], fg=theme["text_primary"], anchor='w').grid(row=row, column=0, sticky='nw', pady=5)
        notes_text_edit = tk.Text(form_frame, font=self.fonts["body"], bg=theme["bg_secondary"], fg=theme["text_primary"], relief='flat', bd=1, height=2, width=30, wrap='word')
        notes_text_edit.grid(row=row, column=1, sticky='ew', pady=5, padx=(10, 0))
        notes_text_edit.insert('1.0', appointment.get('notes', ''))
        edit_vars['notes_widget'] = notes_text_edit
        row += 1

        # Communication preferences (Checkbuttons)
        comm_frame = tk.LabelFrame(form_frame, text="Auto-Reminders", bg=theme["bg_primary"], fg=theme["accent"], font=self.fonts["body"])
        comm_frame.grid(row=row, column=0, columnspan=2, sticky='ew', pady=10)

        enable_reminders_var = tk.BooleanVar(value=appointment.get('enable_reminders', True))
        tk.Checkbutton(comm_frame, text="Enable WhatsApp Reminders", variable=enable_reminders_var, bg=theme["bg_primary"], fg=theme["text_primary"], font=self.fonts["body"]).pack(anchor='w', padx=5, pady=2)
        edit_vars['enable_reminders'] = enable_reminders_var

        enable_email_var = tk.BooleanVar(value=appointment.get('enable_email', True))
        tk.Checkbutton(comm_frame, text="Enable Email Reminders", variable=enable_email_var, bg=theme["bg_primary"], fg=theme["text_primary"], font=self.fonts["body"]).pack(anchor='w', padx=5, pady=2)
        edit_vars['enable_email'] = enable_email_var
        row += 1

        form_frame.grid_columnconfigure(1, weight=1)

        def save_changes():
            edited_email = edit_vars['email'].get().strip()
            if edited_email and not self.validate_email(edited_email):
                self.show_toast("Please enter a valid email address!", "error")
                return
            
            edited_phone = edit_vars['phone_number'].get().strip()
            if not self.clean_phone_number(edited_phone):
                self.show_toast("Please enter a valid WhatsApp phone number with country code!", "error")
                return

            edited_time = edit_vars['appointment_time'].get().strip()
            if not self.validate_time(edited_time):
                self.show_toast("Please enter a valid time (HH:MM format)!", "error")
                return

            edited_clinic_date_obj = edit_vars['clinic_appointment_date'].get_date()
            edited_clinic_date_str = edited_clinic_date_obj.strftime('%Y-%m-%d') if edited_clinic_date_obj else ''

            # Get modality and procedure details to create full procedure string
            edited_modality = edit_vars['modality'].get().strip().upper()
            edited_procedure_details = edit_vars['procedure_details'].get().strip()
            full_procedure = f"{edited_modality}: {edited_procedure_details}" if edited_procedure_details else edited_modality

            # Get health card, handle default selection
            edited_health_card = edit_vars['health_card'].get()
            if edited_health_card == HEALTH_CARD_OPTIONS[0]:
                edited_health_card = ""

            appointment.update({
                'patient_name': edit_vars['patient_name'].get().strip(),
                'procedure': full_procedure, # Combined modality and details
                'phone_number': edited_phone,
                'email': edited_email,
                'health_card': edited_health_card,
                'appointment_date': edit_vars['appointment_date'].get_date().strftime('%Y-%m-%d'),
                'appointment_time': edited_time,
                'clinic_appointment_date': edited_clinic_date_str,
                'enable_reminders': edit_vars['enable_reminders'].get(),
                'enable_email': edit_vars['enable_email'].get(),
                'notes': edit_vars['notes_widget'].get('1.0', tk.END).strip(),
                'updated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })
            self.save_data()
            self.refresh_all()
            dialog.destroy()
            self.show_toast("Appointment updated successfully!", "success")

        tk.Button(dialog, text="💾 Save Changes", command=save_changes, bg=theme["success"], fg="white").pack(pady=10)
        tk.Button(dialog, text="❌ Cancel", command=dialog.destroy, bg=theme["danger"], fg="white").pack(pady=5)

    def delete_appointment(self):
        """Delete selected appointment(s)."""
        selected_items = self.appointments_tree.selection()
        if not selected_items:
            self.show_toast("Please select at least one appointment to delete!", "warning")
            return

        if messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete {len(selected_items)} selected appointment(s)? This action cannot be undone."):
            deleted_count = 0
            # Get IDs of selected items - ensure they're strings for consistent comparison
            ids_to_delete = []
            for item in selected_items:
                item_values = self.appointments_tree.item(item)['values']
                if item_values:  # Make sure values exist
                    ids_to_delete.append(str(item_values[0]))  # Convert to string for comparison
            
            # Filter out appointments that match the IDs to delete
            initial_count = len(self.appointments)
            self.appointments = [apt for apt in self.appointments if str(apt.get('id', '')) not in ids_to_delete]
            deleted_count = initial_count - len(self.appointments)

            if deleted_count > 0:
                self.save_data()
                self.refresh_appointments()
                self.update_stats()  # Add this to update the statistics
                self.show_toast(f"{deleted_count} appointment(s) deleted successfully!", "success")
            else:
                self.show_toast("No appointments were deleted. Please try again.", "warning")

    def update_stats(self):
        """Update sidebar statistics"""
        total = len(self.appointments)
        today_str = date.today().strftime('%Y-%m-%d')
        today_count = sum(1 for apt in self.appointments if apt.get('appointment_date') == today_str)
        reminders_sent = len(self.sent_reminders)
        if hasattr(self, 'stats_labels'):
            self.stats_labels['total'].config(text=str(total))
            self.stats_labels['today'].config(text=str(today_count))
            self.stats_labels['whatsapp_sent'].config(text=str(reminders_sent))
        if self.current_page == "dashboard":
            self.update_dashboard_page()

    def show_settings_page(self):
        """Show settings page and ensure its variables are up to date."""
        self.hide_all_pages()
        self.pages["settings"].pack(fill='both', expand=True)
        self.update_settings_page_vars()

    def export_to_csv(self):
        if not self.appointments: 
            return self.show_toast("No data to export.", "warning")
        filename = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV files", "*.csv")])
        if not filename: 
            return
        try:
            # Dynamically determine fieldnames from all appointments to include new fields
            all_keys = set()
            for apt in self.appointments:
                all_keys.update(apt.keys())
            fieldnames = sorted(list(all_keys)) # Sort for consistent column order

            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                for apt in self.appointments:
                    # Write only fields present in the current appointment, or empty string for missing ones
                    writer.writerow({k: apt.get(k, '') for k in fieldnames})
            self.show_toast("Exported to CSV successfully!", "success")
        except Exception as e:
            self.show_toast(f"Export failed: {e}", "error")

    def export_to_excel(self):
        if not PANDAS_AVAILABLE: 
            return self.show_toast("Pandas library not found.", "error")
        if not self.appointments: 
            return self.show_toast("No data to export.", "warning")
        filename = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if not filename: 
            return
        try:
            pd.DataFrame(self.appointments).to_excel(filename, index=False)
            self.show_toast("Exported to Excel successfully!", "success")
        except Exception as e:
            self.show_toast(f"Export failed: {e}", "error")

    def export_data(self):
        if not self.appointments: 
            return self.show_toast("No data to export.", "warning")
        filename = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON files", "*.json")])
        if not filename: 
            return
        try:
            with open(filename, 'w') as f: 
                json.dump(self.appointments, f, indent=4)
            self.show_toast("Exported to JSON successfully!", "success")
        except Exception as e:
            self.show_toast(f"Export failed: {e}", "error")

    def import_data(self):
        filename = filedialog.askopenfilename(filetypes=[("JSON files", "*.json")])
        if not filename: 
            return
        if not messagebox.askyesno("Confirm Import", "This will overwrite current data. Continue?"): 
            return
        try:
            with open(filename, 'r') as f: 
                raw_data = json.load(f)
                # Validate and clean imported data
                self.appointments = []
                for apt in raw_data:
                    if isinstance(apt, dict):
                        cleaned_apt = self.validate_appointment_data(apt)
                        self.appointments.append(cleaned_apt)
            self.save_data()
            self.refresh_all()
            self.show_toast("Data imported successfully!", "success")
        except Exception as e:
            self.show_toast(f"Import failed: {e}", "error")

    def show_help(self):
        messagebox.showinfo("Help", "Modern Clinic System\n\n- Add, view, edit, and delete patient appointments.\n- Use filters to search for specific records.\n- Export data to various formats from the settings page.\n- Configure auto-reminders for WhatsApp and email.\n\nKeyboard Shortcuts:\nCtrl+S: Save data\nCtrl+N: New patient\nCtrl+R: Refresh\nF1: Help")

    def run(self):
        """Run the application"""
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.root.mainloop()

    def on_closing(self):
        """Handle application closing"""
        if messagebox.askokcancel("Quit", "Do you want to exit? Your data will be saved."):
            self.stop_reminder_system()
            self.save_data()
            self.save_reminder_data()
            self.save_email_settings()
            self.root.destroy()

if __name__ == "__main__":
    app = ModernCompactClinicSystem()
    app.run()
        
